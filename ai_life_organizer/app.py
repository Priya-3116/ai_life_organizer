# =============================================================================
# AI Life Organizer — single-file Flask application
# =============================================================================
# pip install Flask Flask-SQLAlchemy Flask-Login Flask-WTF Flask-Mail \
#             WTForms SQLAlchemy PyMySQL Werkzeug itsdangerous APScheduler \
#             sentence-transformers scikit-learn numpy PyPDF2 python-docx bleach
# =============================================================================

# ── Standard library ──────────────────────────────────────────────────────────
import os
import uuid
import secrets
import hashlib
import json
import logging
import threading
from datetime import datetime, timedelta, date
from functools import wraps

# ── Third-party ───────────────────────────────────────────────────────────────
import bleach
import numpy as np
from flask import (
    Flask, render_template_string, redirect, url_for, flash,
    request, abort, jsonify, send_from_directory,
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user,
    login_required, current_user,
)
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from flask_mail import Mail, Message
from wtforms import (
    StringField, PasswordField, SubmitField, BooleanField,
    TextAreaField, SelectField, DateTimeLocalField, DateField, IntegerField,
)
from wtforms.validators import (
    DataRequired, Email, Length, EqualTo, Optional, NumberRange, ValidationError,
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from apscheduler.schedulers.background import BackgroundScheduler

# ── Lazy AI imports (loaded once at startup) ──────────────────────────────────
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    _AI_AVAILABLE = True
except ImportError:
    _AI_AVAILABLE = False

# ── PDF / DOCX ─────────────────────────────────────────────────────────────────
try:
    import PyPDF2
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

try:
    import docx as python_docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

# V2 feature pack (imported after db is defined)
_v2_ctx = {}

# =============================================================================
# App factory & config
# =============================================================================

def create_app():
    app = Flask(__name__)

    # ── Config ────────────────────────────────────────────────────────────────
    app.config.update(
        SECRET_KEY=os.environ.get("SECRET_KEY", "dev-secret-key-change-in-prod"),
        SQLALCHEMY_DATABASE_URI=os.environ.get(
            "DATABASE_URL",
            "mysql+pymysql://root:password@localhost/ai_life_organizer",
        ),
        SQLALCHEMY_TRACK_MODIFICATIONS=False,
        PERMANENT_SESSION_LIFETIME=timedelta(minutes=30),
        MAX_CONTENT_LENGTH=10 * 1024 * 1024,
        UPLOAD_FOLDER=os.path.join(os.path.dirname(__file__), "uploads"),
        ALLOWED_EXTENSIONS={"pdf", "docx", "jpg", "jpeg", "png"},
        WTF_CSRF_ENABLED=True,
        # Flask-Mail
        MAIL_SERVER=os.environ.get("MAIL_SERVER", "smtp.gmail.com"),
        MAIL_PORT=int(os.environ.get("MAIL_PORT", 587)),
        MAIL_USE_TLS=True,
        MAIL_USERNAME=os.environ.get("MAIL_USERNAME"),
        MAIL_PASSWORD=os.environ.get("MAIL_PASSWORD"),
        MAIL_DEFAULT_SENDER=os.environ.get(
            "MAIL_DEFAULT_SENDER", "noreply@ailifeorganizer.com"
        ),
        DEBUG=True,
    )

    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

    # ── Extensions ────────────────────────────────────────────────────────────
    db.init_app(app)
    login_manager.init_app(app)
    csrf.init_app(app)
    mail.init_app(app)

    login_manager.login_view = "login"
    login_manager.login_message_category = "warning"

    @app.template_filter('weekday_letter')
    def weekday_letter_filter(day_index):
        return ['M','T','W','T','F','S','S'][day_index % 7]

    @app.context_processor
    def _inject_globals():
        from flask_wtf.csrf import generate_csrf
        token = generate_csrf()
        return {
            "csrf_token_field": f'<input type="hidden" name="csrf_token" value="{token}">',
            "csrf_token_value": token,
        }

    # ── Register routes ───────────────────────────────────────────────────────
    _register_routes(app)

    # ── Create tables ─────────────────────────────────────────────────────────
    with app.app_context():
        try:
            db.create_all()
            _migrate_schema(app)
        except Exception as e:
            import logging
            logging.error("Could not connect to database at startup: %s", e)
            logging.error("Make sure MySQL is running. Start it with: net start MySQL80")

    # ── V2 feature context (routes registered after templates load) ───────────
    global _v2_ctx
    _v2_ctx.update({
        "db": db, "mail": mail, "csrf": csrf,
        "sanitize": _sanitize, "get_sentence_model": _get_sentence_model,
        "cosine_similarity": cosine_similarity if _AI_AVAILABLE else None,
        "User": User, "Task": Task, "Goal": Goal, "Note": Note, "MoodLog": MoodLog,
        "Habit": Habit, "HabitCheckin": HabitCheckin, "WeeklyReview": WeeklyReview,
        "NoteLink": NoteLink, "FocusSession": FocusSession, "MoodInsight": MoodInsight,
        "TaskShare": TaskShare, "GoalShare": GoalShare, "TimeLog": TimeLog, "Subtask": Subtask,
    })

    # ── Scheduler ────────────────────────────────────────────────────────────
    if not scheduler.running:
        scheduler.add_job(
            func=_job_trigger_reminders,
            args=[app],
            trigger="interval",
            minutes=5,
            id="trigger_reminders",
        )
        scheduler.add_job(
            func=_job_mark_overdue_goals,
            args=[app],
            trigger="cron",
            hour=0,
            minute=0,
            id="mark_overdue_goals",
        )
        scheduler.start()

    return app

# =============================================================================
# Extensions (module-level singletons)
# =============================================================================
db = SQLAlchemy()
login_manager = LoginManager()
csrf = CSRFProtect()
mail = Mail()
scheduler = BackgroundScheduler(daemon=True)

# AI model — loaded once
_sentence_model = None

def _get_sentence_model():
    global _sentence_model
    if _sentence_model is None and _AI_AVAILABLE:
        try:
            _sentence_model = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as exc:
            logging.warning("Could not load SentenceTransformer: %s", exc)
    return _sentence_model

# =============================================================================
# Models
# =============================================================================

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


class User(UserMixin, db.Model):
    __tablename__ = "users"
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    theme = db.Column(db.String(10), default="dark")
    email_digest = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    tasks = db.relationship("Task", backref="owner", lazy=True, cascade="all, delete-orphan")
    goals = db.relationship("Goal", backref="owner", lazy=True, cascade="all, delete-orphan")
    events = db.relationship("CalendarEvent", backref="owner", lazy=True, cascade="all, delete-orphan")
    notes = db.relationship("Note", backref="owner", lazy=True, cascade="all, delete-orphan")
    documents = db.relationship("Document", backref="owner", lazy=True, cascade="all, delete-orphan")
    reminders = db.relationship("Reminder", backref="owner", lazy=True, cascade="all, delete-orphan")
    reset_tokens = db.relationship("PasswordResetToken", backref="owner", lazy=True, cascade="all, delete-orphan")
    projects = db.relationship("Project", backref="owner", lazy=True, cascade="all, delete-orphan")

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


class Task(db.Model):
    __tablename__ = "tasks"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    priority = db.Column(db.Enum("Low", "Medium", "High", "Critical"), default="Medium", nullable=False)
    status = db.Column(db.Enum("Pending", "In Progress", "Completed"), default="Pending", nullable=False)
    deadline = db.Column(db.DateTime, nullable=True)
    completed_at = db.Column(db.DateTime, nullable=True)
    parent_id = db.Column(db.Integer, db.ForeignKey("tasks.id", ondelete="CASCADE"), nullable=True)
    recurrence = db.Column(db.Enum("None", "Daily", "Weekly"), default="None", nullable=False)
    time_logged_minutes = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    subtasks_rel = db.relationship("Subtask", backref="task", lazy=True, cascade="all, delete-orphan")
    time_logs = db.relationship("TimeLog", backref="task", lazy=True, cascade="all, delete-orphan")


class Goal(db.Model):
    __tablename__ = "goals"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    target_date = db.Column(db.Date, nullable=False)
    progress = db.Column(db.Integer, default=0)
    status = db.Column(db.Enum("Active", "Achieved", "Overdue"), default="Active", nullable=False)
    achieved_at = db.Column(db.Date, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def update_progress(self, value):
        self.progress = max(0, min(100, int(value)))
        if self.progress == 100:
            self.status = "Achieved"
            self.achieved_at = date.today()
        elif self.target_date < date.today() and self.progress < 100:
            self.status = "Overdue"
        else:
            self.status = "Active"


class CalendarEvent(db.Model):
    __tablename__ = "calendar_events"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    start_datetime = db.Column(db.DateTime, nullable=False)
    end_datetime = db.Column(db.DateTime, nullable=True)
    event_type = db.Column(db.Enum("Event", "TaskDeadline", "GoalTarget"), default="Event", nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Note(db.Model):
    __tablename__ = "notes"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text)
    category = db.Column(db.String(100))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


class Document(db.Model):
    __tablename__ = "documents"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    original_name = db.Column(db.String(255), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    extracted_text = db.Column(db.Text)
    parse_failed = db.Column(db.Boolean, default=False)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)


class Reminder(db.Model):
    __tablename__ = "reminders"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    text = db.Column(db.String(500), nullable=False)
    remind_at = db.Column(db.DateTime, nullable=False)
    is_read = db.Column(db.Boolean, default=False)
    source_type = db.Column(db.Enum("Manual", "Task", "Goal", "Event"), default="Manual", nullable=False)
    source_id = db.Column(db.Integer, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class PasswordResetToken(db.Model):
    __tablename__ = "password_reset_tokens"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    token_hash = db.Column(db.String(255), nullable=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    used = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Habit(db.Model):
    __tablename__ = "habits"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    frequency = db.Column(db.Enum("Daily", "Weekly"), default="Daily", nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    streak = db.Column(db.Integer, default=0)
    last_checked = db.Column(db.Date, nullable=True)
    color = db.Column(db.String(20), default="#7c3aed")
    checkins = db.relationship("HabitCheckin", backref="habit", lazy=True, cascade="all, delete-orphan")


class HabitCheckin(db.Model):
    __tablename__ = "habit_checkins"
    id = db.Column(db.Integer, primary_key=True)
    habit_id = db.Column(db.Integer, db.ForeignKey("habits.id", ondelete="CASCADE"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    checked_date = db.Column(db.Date, nullable=False, default=date.today)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Project(db.Model):
    __tablename__ = "projects"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    project_type = db.Column(db.Enum("Personal Project", "Academic Project", "Professional", "Open Source"), default="Personal Project")
    start_date = db.Column(db.String(20))        # e.g. "Aug 2025"
    end_date = db.Column(db.String(20))          # e.g. "Present" or "Dec 2025"
    github_url = db.Column(db.String(500))
    live_url = db.Column(db.String(500))
    description = db.Column(db.Text)             # bullet points, one per line
    technologies = db.Column(db.String(500))     # comma-separated
    achievement = db.Column(db.String(500))      # awards, prizes, etc.
    created_at = db.Column(db.DateTime, default=datetime.utcnow)




class MoodLog(db.Model):
    __tablename__ = 'mood_logs'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id', ondelete='CASCADE'), nullable=False)
    mood = db.Column(db.Integer, nullable=False)  # 1-5
    energy_level = db.Column(db.Integer, nullable=True)
    tags = db.Column(db.String(200), nullable=True)
    note = db.Column(db.String(500))
    logged_date = db.Column(db.Date, nullable=False, default=date.today)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class WeeklyReview(db.Model):
    __tablename__ = "weekly_reviews"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    week_start = db.Column(db.Date, nullable=False)
    week_end = db.Column(db.Date, nullable=False)
    stats_json = db.Column(db.Text)
    narrative = db.Column(db.Text)
    score = db.Column(db.Float, default=0.0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("user_id", "week_start"),)


class NoteLink(db.Model):
    __tablename__ = "note_links"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    source_note_id = db.Column(db.Integer, db.ForeignKey("notes.id", ondelete="CASCADE"), nullable=False)
    target_note_id = db.Column(db.Integer, db.ForeignKey("notes.id", ondelete="CASCADE"), nullable=False)
    similarity = db.Column(db.Float, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("source_note_id", "target_note_id"),)


class FocusSession(db.Model):
    __tablename__ = "focus_sessions"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    task_id = db.Column(db.Integer, db.ForeignKey("tasks.id", ondelete="SET NULL"), nullable=True)
    planned_minutes = db.Column(db.Integer, default=25)
    actual_minutes = db.Column(db.Integer, nullable=False)
    interruptions = db.Column(db.Integer, default=0)
    quality = db.Column(db.Integer, nullable=True)
    session_type = db.Column(db.Enum("Work", "ShortBreak", "LongBreak"), default="Work")
    started_at = db.Column(db.DateTime, nullable=False)
    ended_at = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class MoodInsight(db.Model):
    __tablename__ = "mood_insights"
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    insight_type = db.Column(db.String(50))
    insight_text = db.Column(db.String(500))
    correlation_r = db.Column(db.Float)
    p_value = db.Column(db.Float)
    sample_size = db.Column(db.Integer)
    computed_at = db.Column(db.DateTime, default=datetime.utcnow)


class Subtask(db.Model):
    __tablename__ = "subtasks"
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("tasks.id", ondelete="CASCADE"), nullable=False)
    title = db.Column(db.String(300), nullable=False)
    completed = db.Column(db.Boolean, default=False)
    sort_order = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class TimeLog(db.Model):
    __tablename__ = "time_logs"
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("tasks.id", ondelete="CASCADE"), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    minutes = db.Column(db.Integer, nullable=False)
    note = db.Column(db.String(300))
    logged_at = db.Column(db.DateTime, default=datetime.utcnow)


class TaskShare(db.Model):
    __tablename__ = "task_shares"
    id = db.Column(db.Integer, primary_key=True)
    task_id = db.Column(db.Integer, db.ForeignKey("tasks.id", ondelete="CASCADE"), nullable=False)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("task_id", "shared_with_id"),)


class GoalShare(db.Model):
    __tablename__ = "goal_shares"
    id = db.Column(db.Integer, primary_key=True)
    goal_id = db.Column(db.Integer, db.ForeignKey("goals.id", ondelete="CASCADE"), nullable=False)
    owner_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey("users.id", ondelete="CASCADE"), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    __table_args__ = (db.UniqueConstraint("goal_id", "shared_with_id"),)


def _migrate_schema(app):
    """Add new columns to existing tables (MySQL-safe)."""
    from sqlalchemy import inspect, text
    with app.app_context():
        try:
            insp = inspect(db.engine)
            if "users" in insp.get_table_names():
                cols = {c["name"] for c in insp.get_columns("users")}
                if "theme" not in cols:
                    db.session.execute(text("ALTER TABLE users ADD COLUMN theme VARCHAR(10) DEFAULT 'dark'"))
                if "email_digest" not in cols:
                    db.session.execute(text("ALTER TABLE users ADD COLUMN email_digest BOOLEAN DEFAULT TRUE"))
            if "tasks" in insp.get_table_names():
                cols = {c["name"] for c in insp.get_columns("tasks")}
                if "parent_id" not in cols:
                    db.session.execute(text("ALTER TABLE tasks ADD COLUMN parent_id INT NULL"))
                if "recurrence" not in cols:
                    db.session.execute(text("ALTER TABLE tasks ADD COLUMN recurrence ENUM('None','Daily','Weekly') DEFAULT 'None'"))
                if "time_logged_minutes" not in cols:
                    db.session.execute(text("ALTER TABLE tasks ADD COLUMN time_logged_minutes INT DEFAULT 0"))
            if "mood_logs" in insp.get_table_names():
                cols = {c["name"] for c in insp.get_columns("mood_logs")}
                if "energy_level" not in cols:
                    db.session.execute(text("ALTER TABLE mood_logs ADD COLUMN energy_level INT NULL"))
                if "tags" not in cols:
                    db.session.execute(text("ALTER TABLE mood_logs ADD COLUMN tags VARCHAR(200) NULL"))
            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            logging.warning("Schema migration skipped: %s", exc)


# =============================================================================
# WTForms
# =============================================================================

class RegistrationForm(FlaskForm):
    name = StringField("Name", validators=[DataRequired(), Length(min=1, max=100)])
    email = StringField("Email", validators=[DataRequired(), Email(), Length(max=150)])
    password = PasswordField("Password", validators=[DataRequired(), Length(min=8, message="Password must be at least 8 characters.")])
    confirm_password = PasswordField("Confirm Password", validators=[DataRequired(), EqualTo("password")])
    submit = SubmitField("Register")

    def validate_email(self, field):
        if User.query.filter_by(email=field.data.lower()).first():
            raise ValidationError("Email is already registered.")


class LoginForm(FlaskForm):
    email = StringField("Email", validators=[DataRequired(), Email()])
    password = PasswordField("Password", validators=[DataRequired()])
    remember = BooleanField("Remember Me")
    submit = SubmitField("Log In")


class PasswordResetRequestForm(FlaskForm):
    email = StringField("Email", validators=[DataRequired(), Email()])
    submit = SubmitField("Send Reset Link")


class PasswordResetForm(FlaskForm):
    password = PasswordField("New Password", validators=[DataRequired(), Length(min=8)])
    confirm_password = PasswordField("Confirm Password", validators=[DataRequired(), EqualTo("password")])
    submit = SubmitField("Reset Password")


class ProfileForm(FlaskForm):
    name = StringField("Name", validators=[DataRequired(), Length(min=1, max=100)])
    email = StringField("Email", validators=[DataRequired(), Email(), Length(max=150)])
    theme = SelectField("Theme", choices=[("dark", "Dark"), ("light", "Light")], default="dark")
    email_digest = BooleanField("Weekly email digest (Sundays)")
    submit = SubmitField("Save Changes")


class TaskForm(FlaskForm):
    title = StringField("Title", validators=[DataRequired(), Length(max=200)])
    description = TextAreaField("Description", validators=[Optional(), Length(max=2000)])
    priority = SelectField("Priority", choices=[("Low", "Low"), ("Medium", "Medium"), ("High", "High"), ("Critical", "Critical")], default="Medium")
    status = SelectField("Status", choices=[("Pending", "Pending"), ("In Progress", "In Progress"), ("Completed", "Completed")], default="Pending")
    deadline = DateTimeLocalField("Deadline", format="%Y-%m-%dT%H:%M", validators=[Optional()])
    recurrence = SelectField("Repeat", choices=[("None", "None"), ("Daily", "Daily"), ("Weekly", "Weekly")], default="None")
    submit = SubmitField("Save Task")


class GoalForm(FlaskForm):
    name = StringField("Goal Name", validators=[DataRequired(), Length(max=200)])
    description = TextAreaField("Description", validators=[Optional(), Length(max=2000)])
    target_date = DateField("Target Date", validators=[DataRequired()])
    progress = IntegerField("Progress (%)", validators=[Optional(), NumberRange(min=0, max=100)], default=0)
    submit = SubmitField("Save Goal")

    def validate_target_date(self, field):
        if field.data and field.data <= date.today():
            raise ValidationError("Target date must be in the future.")


class CalendarEventForm(FlaskForm):
    title = StringField("Title", validators=[DataRequired(), Length(max=200)])
    description = TextAreaField("Description", validators=[Optional(), Length(max=2000)])
    start_datetime = DateTimeLocalField("Start", format="%Y-%m-%dT%H:%M", validators=[DataRequired()])
    end_datetime = DateTimeLocalField("End", format="%Y-%m-%dT%H:%M", validators=[Optional()])
    event_type = SelectField("Type", choices=[("Event", "Event"), ("TaskDeadline", "Task Deadline"), ("GoalTarget", "Goal Target")], default="Event")
    submit = SubmitField("Save Event")


class NoteForm(FlaskForm):
    title = StringField("Title", validators=[DataRequired(), Length(max=200)])
    content = TextAreaField("Content", validators=[Optional(), Length(max=10000)])
    category = StringField("Category", validators=[Optional(), Length(max=100)])
    submit = SubmitField("Save Note")


class ReminderForm(FlaskForm):
    text = StringField("Reminder Text", validators=[DataRequired(), Length(max=500)])
    remind_at = DateTimeLocalField("Remind At", format="%Y-%m-%dT%H:%M", validators=[DataRequired()])
    submit = SubmitField("Set Reminder")

    def validate_remind_at(self, field):
        if field.data and field.data <= datetime.utcnow():
            raise ValidationError("Reminder time must be in the future.")


class HabitForm(FlaskForm):
    name = StringField("Habit Name", validators=[DataRequired(), Length(max=200)])
    description = TextAreaField("Description", validators=[Optional()])
    frequency = SelectField("Frequency", choices=[("Daily", "Daily"), ("Weekly", "Weekly")])
    color = StringField("Color", validators=[Optional()], default="#7c3aed")
    submit = SubmitField("Save Habit")


class ProjectForm(FlaskForm):
    title = StringField("Project Title", validators=[DataRequired(), Length(max=200)])
    project_type = SelectField("Type", choices=[
        ("Personal Project", "Personal Project"),
        ("Academic Project", "Academic Project"),
        ("Professional", "Professional"),
        ("Open Source", "Open Source"),
    ])
    start_date = StringField("Start (e.g. Aug 2025)", validators=[Optional(), Length(max=20)])
    end_date = StringField("End (e.g. Present or Dec 2025)", validators=[Optional(), Length(max=20)])
    github_url = StringField("GitHub Link", validators=[Optional(), Length(max=500)])
    live_url = StringField("Live/Demo Link", validators=[Optional(), Length(max=500)])
    description = TextAreaField("Description (one bullet point per line)", validators=[Optional()])
    technologies = StringField("Technologies Used (comma-separated)", validators=[Optional(), Length(max=500)])
    achievement = StringField("Achievement/Award", validators=[Optional(), Length(max=500)])
    submit = SubmitField("Save Project")


# =============================================================================
# Helper utilities
# =============================================================================

def _sanitize(text: str) -> str:
    """Strip all HTML tags from user input."""
    return bleach.clean(text or "", tags=[], strip=True).strip()


def _allowed_file(filename: str, allowed: set) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed


def _upload_dir(user_id: int) -> str:
    from flask import current_app
    path = os.path.join(current_app.config["UPLOAD_FOLDER"], str(user_id))
    os.makedirs(path, exist_ok=True)
    return path


def _extract_text(filepath: str, file_type: str) -> tuple:
    """Return (extracted_text, parse_failed)."""
    text = ""
    failed = False
    try:
        if file_type == "pdf" and _PDF_AVAILABLE:
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages = [page.extract_text() or "" for page in reader.pages]
                text = "\n".join(pages)
        elif file_type == "docx" and _DOCX_AVAILABLE:
            doc = python_docx.Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            text = ""
    except Exception:
        failed = True
    return text, failed


# =============================================================================
# AI Assistant — intent matching
# =============================================================================

_INTENTS = {
    "pending_tasks": [
        "What tasks are pending?",
        "Show me my pending tasks",
        "What do I still need to do?",
        "Tasks not done yet",
    ],
    "upcoming_deadlines": [
        "What deadlines are coming up?",
        "Upcoming due dates",
        "What is due soon?",
        "Deadlines this week",
    ],
    "goal_progress": [
        "How are my goals going?",
        "Show goal progress",
        "Am I on track with my goals?",
        "Goal status update",
    ],
    "focus_today": [
        "What should I focus on today?",
        "Top priority tasks today",
        "What is most important today?",
        "Today's priorities",
    ],
}


def _build_intent_embeddings():
    model = _get_sentence_model()
    if model is None:
        return None, None
    intent_labels = []
    intent_sentences = []
    for intent, sentences in _INTENTS.items():
        for s in sentences:
            intent_labels.append(intent)
            intent_sentences.append(s)
    embeddings = model.encode(intent_sentences, convert_to_numpy=True)
    return intent_labels, embeddings


_intent_labels = None
_intent_embeddings = None


def _match_intent(query: str, threshold: float = 0.45):
    global _intent_labels, _intent_embeddings
    model = _get_sentence_model()
    if model is None:
        return None
    if _intent_labels is None:
        _intent_labels, _intent_embeddings = _build_intent_embeddings()
    if _intent_labels is None:
        return None
    q_emb = model.encode([query], convert_to_numpy=True)
    sims = cosine_similarity(q_emb, _intent_embeddings)[0]
    best_idx = int(np.argmax(sims))
    if sims[best_idx] >= threshold:
        return _intent_labels[best_idx]
    return None


def _handle_intent(intent: str, user_id: int) -> str:
    now = datetime.utcnow()
    if intent == "pending_tasks":
        tasks = Task.query.filter_by(user_id=user_id, status="Pending").order_by(Task.deadline).limit(5).all()
        if not tasks:
            return "You have no pending tasks right now."
        lines = [f"• {t.title} (Priority: {t.priority})" for t in tasks]
        return "Your pending tasks:\n" + "\n".join(lines)

    elif intent == "upcoming_deadlines":
        week = now + timedelta(days=7)
        tasks = (Task.query
                 .filter(Task.user_id == user_id,
                         Task.deadline.isnot(None),
                         Task.deadline <= week,
                         Task.status != "Completed")
                 .order_by(Task.deadline).limit(5).all())
        if not tasks:
            return "No deadlines in the next 7 days."
        lines = [f"• {t.title} — due {t.deadline.strftime('%b %d %H:%M')}" for t in tasks]
        return "Upcoming deadlines:\n" + "\n".join(lines)

    elif intent == "goal_progress":
        goals = Goal.query.filter_by(user_id=user_id).order_by(Goal.target_date).limit(5).all()
        if not goals:
            return "You have no goals set yet."
        lines = [f"• {g.name}: {g.progress}% ({g.status})" for g in goals]
        return "Your goals:\n" + "\n".join(lines)

    elif intent == "focus_today":
        recs = _recommend_tasks(user_id)
        if not recs:
            return "Nothing urgent on your plate today. "
        lines = [f"• {t.title} (Priority: {t.priority})" for t in recs]
        return "Focus on these today:\n" + "\n".join(lines)

    return "I'm not sure how to answer that."


# =============================================================================
# Recommendation Engine
# =============================================================================

_PRIORITY_WEIGHTS = {"Critical": 4, "High": 3, "Medium": 2, "Low": 1}


def _recommend_tasks(user_id: int, top_n: int = 3):
    tasks = Task.query.filter(
        Task.user_id == user_id,
        Task.status != "Completed",
    ).all()

    now = datetime.utcnow()
    active_goals = Goal.query.filter_by(user_id=user_id, status="Active").all()
    avg_gap = (
        sum(100 - g.progress for g in active_goals) / len(active_goals)
        if active_goals else 0
    )

    scored = []
    for t in tasks:
        priority_w = _PRIORITY_WEIGHTS.get(t.priority, 2)
        if t.deadline:
            delta_hours = max(0, (t.deadline - now).total_seconds() / 3600)
            deadline_urgency = max(0, 10 - delta_hours / 24)  # peaks at 0 hours
        else:
            deadline_urgency = 0
        goal_gap = avg_gap / 100  # normalised 0-1
        score = priority_w + deadline_urgency + goal_gap
        scored.append((score, t))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [t for _, t in scored[:top_n]]

# =============================================================================
# Scheduler jobs
# =============================================================================

def _job_trigger_reminders(app):
    """Run every 5 minutes: create advance-warning reminders for upcoming deadlines."""
    with app.app_context():
        try:
            now = datetime.utcnow()

            # ── Auto-create 24h and 1h warnings for tasks ─────────────────
            for hours_before, label in [(24, "24 hours"), (1, "1 hour")]:
                window_start = now + timedelta(hours=hours_before - 0.1)
                window_end   = now + timedelta(hours=hours_before + 0.1)
                tasks = Task.query.filter(
                    Task.deadline >= window_start,
                    Task.deadline <= window_end,
                    Task.status != "Completed",
                ).all()
                for t in tasks:
                    existing = Reminder.query.filter_by(
                        source_type="Task",
                        source_id=t.id,
                        user_id=t.user_id,
                    ).filter(
                        Reminder.text.contains(label)
                    ).first()
                    if not existing:
                        r = Reminder(
                            user_id=t.user_id,
                            text=f"⚠️ Task due in {label}: \"{t.title}\" (Priority: {t.priority})",
                            remind_at=now,
                            source_type="Task",
                            source_id=t.id,
                            is_read=False,
                        )
                        db.session.add(r)

            # ── Auto-create 3-day and 1-day warnings for goals ────────────
            today = date.today()
            for days_before, label in [(3, "3 days"), (1, "tomorrow")]:
                target = today + timedelta(days=days_before)
                goals = Goal.query.filter(
                    Goal.target_date == target,
                    Goal.status == "Active",
                ).all()
                for g in goals:
                    existing = Reminder.query.filter_by(
                        source_type="Goal",
                        source_id=g.id,
                        user_id=g.user_id,
                    ).filter(
                        Reminder.text.contains(label)
                    ).first()
                    if not existing:
                        r = Reminder(
                            user_id=g.user_id,
                            text=f"🎯 Goal deadline in {label}: \"{g.name}\" — {g.progress}% complete",
                            remind_at=now,
                            source_type="Goal",
                            source_id=g.id,
                            is_read=False,
                        )
                        db.session.add(r)

            # ── Warn if a daily habit was not checked in today ────────────
            # Only run this once per day at ~8 AM UTC
            if now.hour == 8:
                habits = Habit.query.filter_by(frequency="Daily").all()
                for h in habits:
                    already_checked = HabitCheckin.query.filter_by(
                        habit_id=h.id,
                        checked_date=today,
                    ).first()
                    if not already_checked:
                        existing = Reminder.query.filter_by(
                            source_type="Manual",
                            user_id=h.user_id,
                        ).filter(
                            Reminder.text.contains(f'"{h.name}"'),
                            Reminder.remind_at >= datetime.combine(today, datetime.min.time()),
                        ).first()
                        if not existing:
                            r = Reminder(
                                user_id=h.user_id,
                                text=f"🔥 Don't break your streak! Check in habit: \"{h.name}\"",
                                remind_at=now,
                                source_type="Manual",
                                source_id=None,
                                is_read=False,
                            )
                            db.session.add(r)

            db.session.commit()
        except Exception as exc:
            db.session.rollback()
            logging.error("Reminder job error: %s", exc)


def _job_mark_overdue_goals(app):
    with app.app_context():
        try:
            today = date.today()
            overdue = Goal.query.filter(
                Goal.target_date < today,
                Goal.status == "Active",
                Goal.progress < 100,
            ).all()
            for g in overdue:
                g.status = "Overdue"
            db.session.commit()
        except Exception as exc:
            logging.error("Overdue goals job error: %s", exc)

# =============================================================================
# Templates
# =============================================================================

BASE_TEMPLATE = """
<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="theme-color" content="#7c3aed">
  <meta name="apple-mobile-web-app-capable" content="yes">
  <link rel="manifest" href="{{ url_for('pwa_manifest') }}">
  <title>{% block title %}AI Life Organizer{% endblock %}</title>
  <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css">
  <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/bootstrap-icons@1.11.3/font/bootstrap-icons.min.css">
  <style>
    :root {
      --bg-primary: #0f1117;
      --bg-secondary: #161b22;
      --bg-card: #161b22;
      --border-color: #30363d;
      --text-primary: #e6edf3;
      --text-secondary: #8b949e;
      --accent-purple: #7c3aed;
      --accent-pink: #ec4899;
      --accent-amber: #f59e0b;
      --accent-green: #10b981;
      --accent-red: #ef4444;
      --accent-teal: #06b6d4;
      --gradient-primary: linear-gradient(135deg, #7c3aed, #ec4899);
    }
    * { box-sizing: border-box; }
    body { background: var(--bg-primary); color: var(--text-primary); font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; margin: 0; }
    a { color: var(--accent-purple); text-decoration: none; }
    a:hover { color: #a78bfa; }
    .sidebar { background: var(--bg-secondary); border-right: 1px solid var(--border-color); width: 250px; min-height: 100vh; position: fixed; top: 0; left: 0; z-index: 100; display: flex; flex-direction: column; }
    .sidebar-logo { padding: 20px 16px 12px; border-bottom: 1px solid var(--border-color); }
    .sidebar-logo .logo-text { font-size: 1.1rem; font-weight: 700; background: var(--gradient-primary); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
    .sidebar-section-label { font-size: 0.65rem; font-weight: 600; color: var(--text-secondary); text-transform: uppercase; letter-spacing: 0.08em; padding: 16px 16px 4px; }
    .sidebar .nav-link { color: var(--text-secondary); padding: 9px 16px; border-radius: 8px; display: flex; align-items: center; gap: 10px; margin: 1px 8px; font-size: 0.88rem; transition: all 0.15s; border-left: 3px solid transparent; }
    .sidebar .nav-link:hover { color: var(--text-primary); background: #ffffff0d; }
    .sidebar .nav-link.active { color: #a78bfa; background: #7c3aed1a; border-left-color: var(--accent-purple); }
    .sidebar .nav-link i { font-size: 1rem; width: 18px; text-align: center; }
    .sidebar-footer { margin-top: auto; padding: 16px; border-top: 1px solid var(--border-color); }
    .sidebar-avatar { width: 32px; height: 32px; border-radius: 50%; background: var(--gradient-primary); display: flex; align-items: center; justify-content: center; font-size: 0.8rem; font-weight: 600; color: white; flex-shrink: 0; }
    .topbar { margin-left: 250px; background: var(--bg-secondary); border-bottom: 1px solid var(--border-color); padding: 10px 24px; display: flex; align-items: center; gap: 12px; position: sticky; top: 0; z-index: 99; }
    .topbar .search-input { background: var(--bg-primary); border: 1px solid var(--border-color); color: var(--text-primary); border-radius: 8px; padding: 6px 12px; font-size: 0.85rem; width: 240px; outline: none; }
    .topbar .search-input::placeholder { color: var(--text-secondary); }
    .topbar .search-input:focus { border-color: var(--accent-purple); }
    .topbar-right { margin-left: auto; display: flex; align-items: center; gap: 12px; }
    .notif-btn { background: none; border: none; color: var(--text-secondary); font-size: 1.2rem; cursor: pointer; position: relative; padding: 4px; transition: color 0.15s; }
    .notif-btn:hover { color: var(--text-primary); }
    .notif-badge { position: absolute; top: -2px; right: -4px; background: var(--accent-red); color: white; font-size: 0.6rem; width: 16px; height: 16px; border-radius: 50%; display: flex; align-items: center; justify-content: center; }
    .user-chip { display: flex; align-items: center; gap: 8px; padding: 4px 10px; border-radius: 8px; background: #ffffff0a; border: 1px solid var(--border-color); cursor: pointer; color: var(--text-primary); font-size: 0.85rem; }
    .main-content { margin-left: 250px; padding: 24px; min-height: 100vh; }
    .card-dark { background: var(--bg-card); border: 1px solid var(--border-color); border-radius: 12px; }
    .card { background: var(--bg-card); border: 1px solid var(--border-color); border-radius: 12px; color: var(--text-primary); }
    .card-header { background: transparent; border-bottom: 1px solid var(--border-color); color: var(--text-primary); font-weight: 600; font-size: 0.9rem; }
    .card-body { color: var(--text-primary); }
    .card-footer { background: transparent; border-top: 1px solid var(--border-color); }
    .stat-card { padding: 20px; border-radius: 12px; }
    .stat-card.purple { background: linear-gradient(135deg, #7c3aed22, #7c3aed44); border: 1px solid #7c3aed55; }
    .stat-card.amber  { background: linear-gradient(135deg, #f59e0b22, #f59e0b44); border: 1px solid #f59e0b55; }
    .stat-card.red    { background: linear-gradient(135deg, #ef444422, #ef444444); border: 1px solid #ef444455; }
    .stat-card.green  { background: linear-gradient(135deg, #10b98122, #10b98144); border: 1px solid #10b98155; }
    .stat-num { font-size: 2rem; font-weight: 700; line-height: 1; display: block; }
    .stat-label { font-size: 0.8rem; color: var(--text-secondary); margin-top: 4px; display: block; }
    .stat-icon { font-size: 2rem; opacity: 0.4; }
    .btn-purple { background: var(--gradient-primary); color: white; border: none; border-radius: 8px; }
    .btn-purple:hover { opacity: 0.9; color: white; }
    .btn-amber { background: var(--accent-amber); color: #0f1117; border: none; border-radius: 8px; }
    .btn-amber:hover { opacity: 0.9; color: #0f1117; }
    .btn-teal { background: var(--accent-teal); color: white; border: none; border-radius: 8px; }
    .btn-teal:hover { opacity: 0.9; color: white; }
    .btn-pink { background: var(--accent-pink); color: white; border: none; border-radius: 8px; }
    .btn-pink:hover { opacity: 0.9; color: white; }
    .btn-green { background: var(--accent-green); color: white; border: none; border-radius: 8px; }
    .btn-green:hover { opacity: 0.9; color: white; }
    .btn-outline-secondary { color: var(--text-secondary); border-color: var(--border-color); background: transparent; }
    .btn-outline-secondary:hover { background: #ffffff11; color: var(--text-primary); border-color: var(--text-secondary); }
    .btn-outline-danger { color: var(--accent-red); border-color: var(--accent-red); background: transparent; }
    .btn-outline-danger:hover { background: #ef444422; color: var(--accent-red); }
    .btn-outline-info { color: var(--accent-teal); border-color: var(--accent-teal); background: transparent; }
    .btn-outline-info:hover { background: #06b6d422; color: var(--accent-teal); }
    .btn-outline-success { color: var(--accent-green); border-color: var(--accent-green); background: transparent; }
    .btn-outline-success:hover { background: #10b98122; color: var(--accent-green); }
    .btn-outline-primary { color: var(--accent-purple); border-color: var(--accent-purple); background: transparent; }
    .btn-outline-primary:hover { background: #7c3aed22; color: var(--accent-purple); }
    .btn-primary { background: var(--accent-purple); border-color: var(--accent-purple); color: white; }
    .btn-primary:hover { background: #6d28d9; border-color: #6d28d9; color: white; }
    .badge-purple { background: #7c3aed33; color: #a78bfa; border: 1px solid #7c3aed55; }
    .badge { border-radius: 6px; font-weight: 500; }
    .bg-secondary { background: #30363d !important; color: var(--text-primary) !important; }
    .bg-success  { background: #10b98133 !important; color: var(--accent-green) !important; border: 1px solid #10b98155; }
    .bg-danger   { background: #ef444433 !important; color: var(--accent-red) !important; border: 1px solid #ef444455; }
    .bg-warning  { background: #f59e0b33 !important; color: var(--accent-amber) !important; border: 1px solid #f59e0b55; }
    .bg-info     { background: #06b6d433 !important; color: var(--accent-teal) !important; border: 1px solid #06b6d455; }
    .bg-primary  { background: #7c3aed33 !important; color: #a78bfa !important; border: 1px solid #7c3aed55; }
    .form-control, .form-select { background: var(--bg-primary); border: 1px solid var(--border-color); color: var(--text-primary); border-radius: 8px; }
    .form-control:focus, .form-select:focus { background: var(--bg-primary); border-color: var(--accent-purple); color: var(--text-primary); box-shadow: 0 0 0 3px #7c3aed22; }
    .form-control::placeholder { color: var(--text-secondary); }
    .form-label { color: var(--text-secondary); font-size: 0.85rem; font-weight: 500; }
    .form-select option { background: var(--bg-secondary); color: var(--text-primary); }
    .form-check-input { background-color: var(--bg-primary); border-color: var(--border-color); }
    .form-check-input:checked { background-color: var(--accent-purple); border-color: var(--accent-purple); }
    .form-check-label { color: var(--text-secondary); }
    .invalid-feedback { color: var(--accent-red); font-size: 0.8rem; }
    .is-invalid { border-color: var(--accent-red) !important; }
    .table { color: var(--text-primary); --bs-table-bg: transparent; --bs-table-striped-bg: transparent; }
    .table > :not(caption) > * > * { background-color: transparent; color: var(--text-primary); border-bottom-color: var(--border-color); }
    .table-light { --bs-table-bg: #ffffff08; color: var(--text-secondary); }
    .table-hover > tbody > tr:hover > * { background-color: #ffffff06; }
    .table thead th { color: var(--text-secondary); font-size: 0.8rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.04em; border-bottom: 1px solid var(--border-color); }
    .list-group-item { background: transparent; border-color: var(--border-color); color: var(--text-primary); }
    .list-group-flush .list-group-item { border-left: 0; border-right: 0; }
    .list-group-item-light { background: #ffffff06 !important; }
    .alert { border-radius: 10px; border: 1px solid; }
    .alert-success { background: #10b98115; border-color: #10b98133; color: #6ee7b7; }
    .alert-danger  { background: #ef444415; border-color: #ef444433; color: #fca5a5; }
    .alert-warning { background: #f59e0b15; border-color: #f59e0b33; color: #fcd34d; }
    .alert-info    { background: #06b6d415; border-color: #06b6d433; color: #67e8f9; }
    .btn-close { filter: invert(1) opacity(0.6); }
    .progress { background: #ffffff11; border-radius: 6px; }
    .progress-bar { border-radius: 6px; }
    .text-muted { color: var(--text-secondary) !important; }
    .text-primary { color: var(--accent-purple) !important; }
    .text-success { color: var(--accent-green) !important; }
    .text-danger  { color: var(--accent-red) !important; }
    .text-warning { color: var(--accent-amber) !important; }
    .text-info    { color: var(--accent-teal) !important; }
    .priority-Critical { color: var(--accent-red); font-weight: 600; }
    .priority-High   { color: var(--accent-amber); }
    .priority-Medium { color: #a78bfa; }
    .priority-Low    { color: var(--text-secondary); }
    .dashboard-hero { background: linear-gradient(135deg, #7c3aed1a, #ec489911); border: 1px solid #7c3aed33; border-radius: 16px; padding: 28px 32px; margin-bottom: 24px; display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 20px; }
    .dashboard-hero h1 { font-size: 1.6rem; font-weight: 700; margin: 0 0 4px; }
    .hero-stats { display: flex; gap: 32px; }
    .hero-stat { text-align: center; }
    .hero-stat .stat-num { font-size: 1.8rem; font-weight: 700; background: var(--gradient-primary); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
    .hero-stat span:last-child { font-size: 0.75rem; color: var(--text-secondary); display: block; }
    .quick-create { display: flex; flex-wrap: wrap; gap: 8px; padding: 16px 20px; background: var(--bg-card); border: 1px solid var(--border-color); border-radius: 12px; align-items: center; }
    .quick-create span { font-size: 0.8rem; color: var(--text-secondary); margin-right: 4px; }
    .gauge-container { position: relative; width: 160px; height: 160px; margin: 0 auto; }
    .gauge-score { position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); text-align: center; }
    .gauge-score .score-num { font-size: 2.2rem; font-weight: 700; line-height: 1; }
    .gauge-score .score-label { font-size: 0.7rem; color: var(--text-secondary); }
    .clock-widget { text-align: center; padding: 20px; }
    .clock-time { font-size: 2rem; font-weight: 700; font-family: 'Courier New', monospace; background: var(--gradient-primary); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; }
    .clock-date { font-size: 0.78rem; color: var(--text-secondary); margin-top: 4px; }
    .timeline-item { display: flex; gap: 12px; padding: 10px 0; border-bottom: 1px solid var(--border-color); }
    .timeline-item:last-child { border-bottom: 0; }
    .timeline-dot { width: 8px; height: 8px; border-radius: 50%; margin-top: 6px; flex-shrink: 0; }
    .habit-card { background: var(--bg-card); border: 1px solid var(--border-color); border-radius: 12px; padding: 16px; }
    .habit-grid { display: flex; gap: 4px; margin-top: 8px; }
    .habit-dot { width: 18px; height: 18px; border-radius: 4px; background: var(--border-color); }
    .habit-dot.done { background: var(--accent-green); }
    .streak-badge { display: inline-flex; align-items: center; gap: 4px; font-size: 0.8rem; font-weight: 600; color: var(--accent-amber); }
    .dropdown-menu { background: var(--bg-secondary); border: 1px solid var(--border-color); border-radius: 10px; }
    .dropdown-item { color: var(--text-secondary); font-size: 0.88rem; }
    .dropdown-item:hover { background: #ffffff0d; color: var(--text-primary); }
    .dropdown-divider { border-color: var(--border-color); }
    .dropdown-item.text-danger { color: var(--accent-red) !important; }
    ::-webkit-scrollbar { width: 6px; height: 6px; }
    ::-webkit-scrollbar-track { background: transparent; }
    ::-webkit-scrollbar-thumb { background: var(--border-color); border-radius: 3px; }
    ::-webkit-scrollbar-thumb:hover { background: var(--text-secondary); }
    @media (max-width: 768px) {
      .sidebar { transform: translateX(-100%); transition: transform 0.25s; }
      .sidebar.open { transform: translateX(0); }
      .main-content, .topbar { margin-left: 0; }
      .dashboard-hero { flex-direction: column; }
      .hero-stats { gap: 16px; }
    }
    body.theme-light {
      --bg-primary: #f6f8fa; --bg-secondary: #ffffff; --bg-card: #ffffff;
      --border-color: #d0d7de; --text-primary: #1f2328; --text-secondary: #656d76;
    }
    body.theme-light .sidebar { background: #ffffff; }
    body.theme-light .topbar { background: #ffffff; }
    body.theme-light .card { background: #ffffff; }
  </style>
  {% block extra_css %}{% endblock %}
</head>
<body class="{% if current_user.is_authenticated %}theme-{{ current_user.theme or 'dark' }}{% endif %}">
{% if current_user.is_authenticated %}
<div class="sidebar" id="sidebar">
  <div class="sidebar-logo">
    <a href="{{ url_for('dashboard') }}" class="d-flex align-items-center gap-2" style="text-decoration:none">
      <i class="bi bi-grid-3x3-gap" style="font-size:1.3rem;color:var(--accent-purple)"></i>
      <span class="logo-text">Life OS</span>
    </a>
  </div>
  <div class="sidebar-section-label">Main</div>
  <nav>
    <a href="{{ url_for('dashboard') }}"      class="nav-link {% if request.endpoint == 'dashboard' %}active{% endif %}"><i class="bi bi-grid-1x2"></i> Dashboard</a>
    <a href="{{ url_for('tasks_index') }}"    class="nav-link {% if request.endpoint and 'task' in request.endpoint %}active{% endif %}"><i class="bi bi-check2-square"></i> Tasks</a>
    <a href="{{ url_for('goals_index') }}"    class="nav-link {% if request.endpoint and 'goal' in request.endpoint %}active{% endif %}"><i class="bi bi-trophy"></i> Goals</a>
    <a href="{{ url_for('calendar_index') }}" class="nav-link {% if request.endpoint and 'calendar' in request.endpoint %}active{% endif %}"><i class="bi bi-calendar3"></i> Calendar</a>
    <a href="{{ url_for('notes_index') }}"    class="nav-link {% if request.endpoint and 'note' in request.endpoint %}active{% endif %}"><i class="bi bi-journal-text"></i> Notes</a>
    <a href="{{ url_for('documents_index') }}"class="nav-link {% if request.endpoint and 'document' in request.endpoint %}active{% endif %}"><i class="bi bi-folder2"></i> Documents</a>
    <a href="{{ url_for('ai_chat') }}"        class="nav-link {% if request.endpoint == 'ai_chat' %}active{% endif %}"><i class="bi bi-robot"></i> AI Assistant</a>
    <a href="{{ url_for('analytics_index') }}"class="nav-link {% if request.endpoint == 'analytics_index' %}active{% endif %}"><i class="bi bi-bar-chart-line"></i> Analytics</a>
    <a href="{{ url_for('projects_index') }}" class="nav-link {% if request.endpoint and 'project' in request.endpoint %}active{% endif %}"><i class="bi bi-briefcase"></i> Projects</a>
  </nav>
  <div class="sidebar-section-label">Intelligence</div>
  <nav>
    <a href="{{ url_for('weekly_review') }}"    class="nav-link {% if request.endpoint == 'weekly_review' %}active{% endif %}"><i class="bi bi-calendar-week"></i> Weekly Review</a>
    <a href="{{ url_for('focus_index') }}"      class="nav-link {% if request.endpoint == 'focus_index' %}active{% endif %}"><i class="bi bi-stopwatch"></i> Focus Timer</a>
    <a href="{{ url_for('mood_intelligence') }}"class="nav-link {% if request.endpoint == 'mood_intelligence' %}active{% endif %}"><i class="bi bi-emoji-smile"></i> Mood Intel</a>
    <a href="{{ url_for('notes_graph') }}"      class="nav-link {% if request.endpoint == 'notes_graph' %}active{% endif %}"><i class="bi bi-diagram-3"></i> Knowledge Graph</a>
    <a href="{{ url_for('tasks_kanban') }}"     class="nav-link {% if request.endpoint == 'tasks_kanban' %}active{% endif %}"><i class="bi bi-kanban"></i> Kanban</a>
  </nav>
  <div class="sidebar-section-label">Tools</div>
  <nav>
    <a href="{{ url_for('habits_index') }}"   class="nav-link {% if request.endpoint and 'habit' in request.endpoint %}active{% endif %}"><i class="bi bi-activity"></i> Habits</a>
    <a href="{{ url_for('reminders_panel') }}"class="nav-link {% if request.endpoint and 'reminder' in request.endpoint %}active{% endif %}"><i class="bi bi-bell"></i> Reminders</a>
    <a href="{{ url_for('export_import_page') }}" class="nav-link {% if request.endpoint == 'export_import_page' %}active{% endif %}"><i class="bi bi-arrow-down-up"></i> Export/Import</a>
    <a href="{{ url_for('profile') }}"        class="nav-link {% if request.endpoint == 'profile' %}active{% endif %}"><i class="bi bi-person"></i> Profile</a>
  </nav>
  <div class="sidebar-footer">
    <div class="d-flex align-items-center gap-2">
      <div class="sidebar-avatar">{{ current_user.name[0].upper() }}</div>
      <div>
        <div style="font-size:0.82rem;font-weight:600">{{ current_user.name }}</div>
        <a href="{{ url_for('logout') }}" style="font-size:0.72rem;color:var(--text-secondary)"><i class="bi bi-box-arrow-right"></i> Logout</a>
      </div>
    </div>
  </div>
</div>
<div class="topbar">
  <i class="bi bi-search" style="color:var(--text-secondary);font-size:0.9rem"></i>
  <input type="text" class="search-input" placeholder="Search..." id="global-search">
  <div class="topbar-right">
    <form method="POST" action="{{ url_for('set_theme') }}" class="d-flex align-items-center">
      {{ csrf_token_field|default('', true)|safe }}
      <input type="hidden" name="theme" value="{{ 'light' if (current_user.theme or 'dark') == 'dark' else 'dark' }}">
      <button type="submit" class="btn btn-sm btn-outline-secondary" title="Toggle theme"><i class="bi bi-{{ 'sun' if (current_user.theme or 'dark') == 'dark' else 'moon' }}"></i></button>
    </form>
    <a href="{{ url_for('reminders_panel') }}" style="text-decoration:none">
      <button class="notif-btn" title="Reminders">
        <i class="bi bi-bell"></i>
        <span id="notif-badge" class="notif-badge" style="display:none">0</span>
      </button>
    </a>
    <div class="dropdown">
      <div class="user-chip dropdown-toggle" data-bs-toggle="dropdown">
        <div class="sidebar-avatar" style="width:24px;height:24px;font-size:0.65rem">{{ current_user.name[0].upper() }}</div>
        <span>{{ current_user.name }}</span>
      </div>
      <ul class="dropdown-menu dropdown-menu-end">
        <li><a class="dropdown-item" href="{{ url_for('profile') }}"><i class="bi bi-person me-2"></i>Profile</a></li>
        <li><hr class="dropdown-divider"></li>
        <li><a class="dropdown-item text-danger" href="{{ url_for('logout') }}"><i class="bi bi-box-arrow-right me-2"></i>Logout</a></li>
      </ul>
    </div>
  </div>
</div>
{% endif %}
<div class="{% if current_user.is_authenticated %}main-content{% else %}container mt-4{% endif %}">
  {% with messages = get_flashed_messages(with_categories=true) %}
    {% if messages %}
      {% for category, message in messages %}
        <div class="alert alert-{{ category }} alert-dismissible fade show" role="alert">
          {{ message }}
          <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
      {% endfor %}
    {% endif %}
  {% endwith %}
  {% block content %}{% endblock %}
</div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
{% if current_user.is_authenticated %}
<script>
function updateNotifBadge() {
  fetch("{{ url_for('reminders_unread_count') }}")
    .then(r => r.json())
    .then(d => {
      const b = document.getElementById("notif-badge");
      if (d.count > 0) { b.textContent = d.count; b.style.display = "flex"; }
      else { b.style.display = "none"; }
    }).catch(() => {});
}
updateNotifBadge();
setInterval(updateNotifBadge, 60000);
if('serviceWorker' in navigator){ navigator.serviceWorker.register("{{ url_for('service_worker') }}").catch(()=>{}); }
</script>
{% endif %}
{% block extra_js %}{% endblock %}
</body>
</html>
"""

REGISTER_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-6 col-lg-5">
    <div class="card shadow-sm mt-4">
      <div class="card-body p-4">
        <h4 class="card-title mb-4"><i class="bi bi-person-plus"></i> Create Account</h4>
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.name.label(class="form-label") }}
            {{ form.name(class="form-control" + (" is-invalid" if form.name.errors else "")) }}
            {% for e in form.name.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.email.label(class="form-label") }}
            {{ form.email(class="form-control" + (" is-invalid" if form.email.errors else "")) }}
            {% for e in form.email.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.password.label(class="form-label") }}
            {{ form.password(class="form-control" + (" is-invalid" if form.password.errors else "")) }}
            {% for e in form.password.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.confirm_password.label(class="form-label") }}
            {{ form.confirm_password(class="form-control" + (" is-invalid" if form.confirm_password.errors else "")) }}
            {% for e in form.confirm_password.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          {{ form.submit(class="btn btn-primary w-100") }}
        </form>
        <hr>
        <p class="text-center mb-0">Already have an account? <a href="{{ url_for('login') }}">Log in</a></p>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

LOGIN_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-5 col-lg-4">
    <div class="card shadow-sm mt-4">
      <div class="card-body p-4">
        <h4 class="card-title mb-4"><i class="bi bi-box-arrow-in-right"></i> Log In</h4>
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.email.label(class="form-label") }}
            {{ form.email(class="form-control" + (" is-invalid" if form.email.errors else "")) }}
            {% for e in form.email.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.password.label(class="form-label") }}
            {{ form.password(class="form-control" + (" is-invalid" if form.password.errors else "")) }}
            {% for e in form.password.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3 form-check">
            {{ form.remember(class="form-check-input") }}
            {{ form.remember.label(class="form-check-label") }}
          </div>
          {{ form.submit(class="btn btn-primary w-100") }}
        </form>
        <hr>
        <p class="text-center mb-1"><a href="{{ url_for('reset_request') }}">Forgot password?</a></p>
        <p class="text-center mb-0">No account? <a href="{{ url_for('register') }}">Register</a></p>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

RESET_REQUEST_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-5">
    <div class="card shadow-sm mt-4">
      <div class="card-body p-4">
        <h4 class="mb-4"><i class="bi bi-key"></i> Reset Password</h4>
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.email.label(class="form-label") }}
            {{ form.email(class="form-control" + (" is-invalid" if form.email.errors else "")) }}
            {% for e in form.email.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          {{ form.submit(class="btn btn-primary w-100") }}
        </form>
        <hr>
        <p class="text-center mb-0"><a href="{{ url_for('login') }}">Back to login</a></p>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

RESET_PASSWORD_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-5">
    <div class="card shadow-sm mt-4">
      <div class="card-body p-4">
        <h4 class="mb-4"><i class="bi bi-shield-lock"></i> Set New Password</h4>
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.password.label(class="form-label") }}
            {{ form.password(class="form-control" + (" is-invalid" if form.password.errors else "")) }}
            {% for e in form.password.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.confirm_password.label(class="form-label") }}
            {{ form.confirm_password(class="form-control" + (" is-invalid" if form.confirm_password.errors else "")) }}
            {% for e in form.confirm_password.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          {{ form.submit(class="btn btn-primary w-100") }}
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

PROFILE_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-6">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0"><i class="bi bi-person-gear"></i> My Profile</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.name.label(class="form-label") }}
            {{ form.name(class="form-control" + (" is-invalid" if form.name.errors else "")) }}
            {% for e in form.name.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.email.label(class="form-label") }}
            {{ form.email(class="form-control" + (" is-invalid" if form.email.errors else "")) }}
            {% for e in form.email.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="row g-3 mb-3">
            <div class="col-md-6">{{ form.theme.label(class="form-label") }}{{ form.theme(class="form-select") }}</div>
            <div class="col-md-6 d-flex align-items-end">{{ form.email_digest(class="form-check-input me-2") }}{{ form.email_digest.label(class="form-check-label") }}</div>
          </div>
          {{ form.submit(class="btn btn-primary") }}
        </form>
        <hr>
        <p class="text-muted small">Member since {{ current_user.created_at.strftime('%B %Y') }}</p>
        <form method="POST" action="{{ url_for('delete_account') }}" onsubmit="return confirm('Permanently delete your account and all data?')">
          {{ form.hidden_tag() }}
          <button class="btn btn-outline-danger btn-sm">Delete Account</button>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── Dashboard template ──────────────────────────────────────────────────────

DASHBOARD_TEMPLATE = (
    BASE_TEMPLATE
    .replace("{% block content %}{% endblock %}", """{% block content %}
<!-- Hero -->
<div class="dashboard-hero">
  <div>
    <h1>Welcome Back, {{ current_user.name }}! &#128075;</h1>
    <p class="text-muted mb-0">{{ greeting }} &mdash; Here's your life overview</p>
  </div>
  <div class="hero-stats">
    <div class="hero-stat">
      <span class="stat-num">{{ stats.total_tasks }}</span>
      <span>Total Tasks</span>
    </div>
    <div class="hero-stat">
      <span class="stat-num">{{ stats.pending }}</span>
      <span>Pending</span>
    </div>
    <div class="hero-stat">
      <span class="stat-num">{{ stats.goals_active }}</span>
      <span>Active Goals</span>
    </div>
  </div>
</div>

<!-- Quick Create -->
<div class="quick-create mb-4">
  <span>Quick Add:</span>
  <a href="{{ url_for('tasks_new') }}" class="btn btn-sm btn-purple">+ New Task</a>
  <a href="{{ url_for('goals_new') }}" class="btn btn-sm btn-amber">+ New Goal</a>
  <a href="{{ url_for('calendar_new') }}" class="btn btn-sm btn-teal">+ New Event</a>
  <a href="{{ url_for('notes_new') }}" class="btn btn-sm btn-pink">+ New Note</a>
  <a href="{{ url_for('habits_index') }}" class="btn btn-sm btn-green">&#10003; Habits</a>
  <a href="{{ url_for('weekly_review') }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-calendar-week"></i> Weekly Review</a>
</div>

<!-- NL Quick Capture -->
<div class="card mb-4 border-start border-4" style="border-color:var(--accent-teal)!important">
  <div class="card-body py-2">
    <div class="input-group">
      <span class="input-group-text bg-transparent border-secondary"><i class="bi bi-magic"></i></span>
      <input type="text" id="dash-nl" class="form-control" placeholder='Quick capture: "Call dentist tomorrow, high priority"'>
      <a id="dash-nl-go" href="{{ url_for('tasks_new') }}" class="btn btn-purple" style="display:none">Create</a>
    </div>
  </div>
</div>

<!-- Row 1: 4 stat cards -->
<div class="row g-3 mb-3">
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card purple">
      <div class="d-flex justify-content-between align-items-start">
        <div>
          <span class="stat-num" style="color:#a78bfa">{{ stats.pending }}</span>
          <span class="stat-label">Pending Tasks</span>
        </div>
        <i class="bi bi-check2-square stat-icon" style="color:#7c3aed"></i>
      </div>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card amber">
      <div class="d-flex justify-content-between align-items-start">
        <div>
          <span class="stat-num" style="color:#fbbf24">{{ stats.due_today }}</span>
          <span class="stat-label">Due Today</span>
        </div>
        <i class="bi bi-calendar-event stat-icon" style="color:#f59e0b"></i>
      </div>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card red">
      <div class="d-flex justify-content-between align-items-start">
        <div>
          <span class="stat-num" style="color:#f87171">{{ stats.overdue }}</span>
          <span class="stat-label">Overdue Tasks</span>
        </div>
        <i class="bi bi-exclamation-triangle stat-icon" style="color:#ef4444"></i>
      </div>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card green">
      <div class="d-flex justify-content-between align-items-start">
        <div>
          <span class="stat-num" style="color:#34d399">{{ stats.goals_active }}</span>
          <span class="stat-label">Active Goals</span>
        </div>
        <i class="bi bi-trophy stat-icon" style="color:#10b981"></i>
      </div>
    </div>
  </div>
</div>

<!-- Row 2: Weekly Activity + AI Score -->
<div class="row g-3 mb-3">
  <div class="col-lg-8">
    <div class="card h-100">
      <div class="card-header d-flex align-items-center gap-2">
        <i class="bi bi-bar-chart-fill" style="color:var(--accent-purple)"></i>
        Weekly Activity
      </div>
      <div class="card-body">
        <canvas id="weeklyActivityChart" style="max-height:200px"></canvas>
      </div>
    </div>
  </div>
  <div class="col-lg-4">
    <div class="card h-100">
      <div class="card-header d-flex align-items-center gap-2">
        <i class="bi bi-speedometer2" style="color:var(--accent-amber)"></i>
        AI Score
      </div>
      <div class="card-body d-flex flex-column align-items-center justify-content-center gap-3">
        <div class="gauge-container">
          <svg width="160" height="160" viewBox="0 0 160 160">
            <circle cx="80" cy="80" r="60" fill="none" stroke="#30363d" stroke-width="12"/>
            <circle cx="80" cy="80" r="60" fill="none"
              stroke="{{ '#10b981' if ai_score > 70 else ('#f59e0b' if ai_score >= 40 else '#ef4444') }}"
              stroke-width="12"
              stroke-dasharray="{{ (ai_score / 100 * 376.99)|round(1) }} 376.99"
              stroke-dashoffset="94.25"
              stroke-linecap="round"
              transform="rotate(-90 80 80)"/>
          </svg>
          <div class="gauge-score">
            <span class="score-num" style="color:{{ '#10b981' if ai_score > 70 else ('#f59e0b' if ai_score >= 40 else '#ef4444') }}">{{ ai_score }}</span>
            <span class="score-label">/ 100</span>
          </div>
        </div>
        <div class="text-center">
          <div style="font-size:0.85rem;font-weight:600;color:{{ '#10b981' if ai_score > 70 else ('#f59e0b' if ai_score >= 40 else '#ef4444') }}">
            {{ 'Excellent' if ai_score > 70 else ('Good' if ai_score >= 40 else 'Needs Work') }}
          </div>
          <div class="clock-widget">
            <div class="clock-time" id="live-clock">--:--:--</div>
            <div class="clock-date" id="live-date"></div>
          </div>
        </div>
      </div>
    </div>
  </div>
</div>

<!-- Row 3: Upcoming Deadlines + AI Recommendations -->
<div class="row g-3 mb-3">
  <div class="col-lg-6">
    <div class="card h-100">
      <div class="card-header d-flex justify-content-between align-items-center">
        <span><i class="bi bi-calendar-week me-1" style="color:var(--accent-amber)"></i>Upcoming Deadlines</span>
        <a href="{{ url_for('calendar_index') }}" class="btn btn-sm btn-outline-secondary" style="font-size:0.75rem">View All</a>
      </div>
      <div class="card-body py-2">
        {% for t in upcoming_deadlines %}
        <div class="timeline-item">
          <div class="timeline-dot" style="background:{{ '#ef4444' if t.priority == 'Critical' else ('#f59e0b' if t.priority == 'High' else '#7c3aed') }}"></div>
          <div class="flex-grow-1">
            <div style="font-size:0.88rem;font-weight:500">{{ t.title }}</div>
            <div style="font-size:0.75rem;color:var(--text-secondary)">{{ t.deadline.strftime('%b %d, %Y %H:%M') if t.deadline else 'No deadline' }}</div>
          </div>
          <span class="badge bg-{{ 'danger' if t.priority == 'Critical' else ('warning' if t.priority == 'High' else 'secondary') }}" style="font-size:0.7rem;align-self:flex-start">{{ t.priority }}</span>
        </div>
        {% else %}
        <p class="text-muted small mt-2">No upcoming deadlines.</p>
        {% endfor %}
      </div>
    </div>
  </div>
  <div class="col-lg-6">
    <div class="card h-100">
      <div class="card-header">
        <i class="bi bi-lightbulb me-1" style="color:var(--accent-amber)"></i>AI Recommendations
      </div>
      <div class="card-body py-2">
        {% for t in recommendations %}
        <div class="timeline-item">
          <div class="timeline-dot" style="background:var(--accent-purple)"></div>
          <div class="flex-grow-1">
            <div style="font-size:0.88rem;font-weight:500">{{ t.title }}</div>
            <div style="font-size:0.75rem;color:var(--text-secondary)">Priority: {{ t.priority }}</div>
          </div>
          <a href="{{ url_for('tasks_edit', task_id=t.id) }}" style="font-size:0.75rem;color:var(--accent-purple)">View</a>
        </div>
        {% else %}
        <p class="text-muted small mt-2">All caught up! No urgent tasks.</p>
        {% endfor %}
        {% if habits %}
        <div style="margin-top:12px;padding-top:12px;border-top:1px solid var(--border-color)">
          <div style="font-size:0.8rem;color:var(--text-secondary);margin-bottom:8px">Today's Habits</div>
          {% for h in habits[:4] %}
          <div class="d-flex align-items-center justify-content-between mb-2">
            <span style="font-size:0.85rem">{{ h.name }}</span>
            {% if h.id in checked_today %}
              <span class="badge bg-success" style="font-size:0.7rem">Done &#10003;</span>
            {% else %}
              <a href="{{ url_for('habits_checkin', habit_id=h.id) }}" class="badge bg-secondary" style="font-size:0.7rem;cursor:pointer;text-decoration:none" onclick="event.preventDefault();fetch('{{ url_for('habits_checkin', habit_id=h.id) }}',{method:'POST',headers:{'X-CSRFToken':'{{ csrf_token() }}','Content-Type':'application/x-www-form-urlencoded'},body:'csrf_token={{ csrf_token() }}'}).then(()=>location.reload())">Check In</a>
            {% endif %}
          </div>
          {% endfor %}
        </div>
        {% endif %}
      </div>
    </div>
  </div>
</div>

<!-- Row 4: Goal Progress + Recent Notes -->
<div class="row g-3">
  <div class="col-lg-6">
    <div class="card">
      <div class="card-header d-flex justify-content-between align-items-center">
        <span><i class="bi bi-trophy me-1" style="color:var(--accent-green)"></i>Goal Progress</span>
        <a href="{{ url_for('goals_index') }}" class="btn btn-sm btn-outline-secondary" style="font-size:0.75rem">All Goals</a>
      </div>
      <div class="card-body">
        {% for g in goals %}
        <div class="mb-3">
          <div class="d-flex justify-content-between mb-1">
            <span style="font-size:0.88rem">{{ g.name }}</span>
            <span style="font-size:0.8rem;color:{{ '#10b981' if g.status == 'Achieved' else ('#ef4444' if g.status == 'Overdue' else 'var(--accent-purple)') }}">{{ g.progress }}%</span>
          </div>
          <div class="progress" style="height:5px">
            <div class="progress-bar" style="width:{{ g.progress }}%;background:{{ '#10b981' if g.status == 'Achieved' else ('#ef4444' if g.status == 'Overdue' else 'var(--accent-purple)') }}"></div>
          </div>
        </div>
        {% else %}
        <p class="text-muted small">No goals yet.</p>
        {% endfor %}
      </div>
    </div>
  </div>
  <div class="col-lg-6">
    <div class="card">
      <div class="card-header d-flex justify-content-between align-items-center">
        <span><i class="bi bi-journal-text me-1" style="color:var(--accent-teal)"></i>Recent Notes</span>
        <a href="{{ url_for('notes_index') }}" class="btn btn-sm btn-outline-secondary" style="font-size:0.75rem">All Notes</a>
      </div>
      <div class="card-body p-0">
        <ul class="list-group list-group-flush">
          {% for n in recent_notes %}
          <li class="list-group-item">
            <div class="d-flex justify-content-between align-items-start">
              <div>
                <div style="font-size:0.88rem;font-weight:500">{{ n.title }}</div>
                <div style="font-size:0.75rem;color:var(--text-secondary)">{{ n.updated_at.strftime('%b %d, %Y') }}</div>
              </div>
              {% if n.category %}<span class="badge badge-purple" style="font-size:0.7rem">{{ n.category }}</span>{% endif %}
            </div>
          </li>
          {% else %}
          <li class="list-group-item text-muted" style="font-size:0.88rem">No notes yet.</li>
          {% endfor %}
        </ul>
      </div>
    </div>
  </div>
</div>
{% endblock %}""")
    .replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.3/dist/chart.umd.min.js"></script>
<script>
// Live clock
function updateClock() {
  const now = new Date();
  const cl = document.getElementById('live-clock');
  const dt = document.getElementById('live-date');
  if(cl) cl.textContent = now.toLocaleTimeString('en-US', {hour:'2-digit', minute:'2-digit', second:'2-digit'});
  if(dt) dt.textContent = now.toLocaleDateString('en-US', {weekday:'long', year:'numeric', month:'long', day:'numeric'});
}
setInterval(updateClock, 1000);
updateClock();

// Weekly activity chart
const wLabels = {{ weekly_labels|safe }};
const wValues = {{ weekly_values|safe }};
const colors = wValues.map((v, i) => {
  const palette = ['#7c3aed','#ec4899','#f59e0b','#10b981','#06b6d4','#8b5cf6','#f43f5e'];
  return palette[i % palette.length];
});
new Chart(document.getElementById('weeklyActivityChart'), {
  type: 'bar',
  data: {
    labels: wLabels,
    datasets: [{
      label: 'Completed Tasks',
      data: wValues,
      backgroundColor: colors,
      borderRadius: 6,
      borderSkipped: false,
    }]
  },
  options: {
    responsive: true,
    maintainAspectRatio: false,
    plugins: { legend: { display: false } },
    scales: {
      x: { grid: { color: '#30363d' }, ticks: { color: '#8b949e' } },
      y: { beginAtZero: true, ticks: { stepSize: 1, color: '#8b949e' }, grid: { color: '#30363d' } }
    }
  }
});
const dashNl = document.getElementById('dash-nl');
const dashGo = document.getElementById('dash-nl-go');
if(dashNl){
  let dt;
  dashNl.addEventListener('input', ()=>{
    clearTimeout(dt);
    dt = setTimeout(()=>{
      const text = dashNl.value.trim();
      if(!text){ dashGo.style.display='none'; return; }
      fetch("{{ url_for('tasks_parse_nl') }}",{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({text})})
      .then(r=>r.json()).then(d=>{
        dashGo.style.display='inline-block';
        let url = "{{ url_for('tasks_new') }}?title="+encodeURIComponent(d.title)+"&priority="+encodeURIComponent(d.priority);
        if(d.deadline) url += "&deadline="+encodeURIComponent(d.deadline.slice(0,16));
        dashGo.href = url;
      });
    }, 400);
  });
}
</script>
{% endblock %}""")
)

TASKS_INDEX_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-check2-square"></i> Tasks</h3>
  <div class="d-flex gap-2">
    <a href="{{ url_for('tasks_kanban') }}" class="btn btn-outline-secondary btn-sm"><i class="bi bi-kanban"></i> Kanban</a>
    <a href="{{ url_for('tasks_new') }}" class="btn btn-primary"><i class="bi bi-plus-lg"></i> New Task</a>
  </div>
</div>

<!-- NL Task Capture -->
<div class="card mb-4 border-start border-4" style="border-color:var(--accent-purple)!important">
  <div class="card-body">
    <label class="form-label"><i class="bi bi-magic"></i> Natural Language Capture</label>
    <input type="text" id="nl-input" class="form-control mb-2" placeholder='e.g. "Finish report by Friday, high priority"'>
    <div id="nl-preview" style="display:none" class="card card-dark mt-2">
      <div class="card-body py-2">
        <div class="d-flex justify-content-between align-items-start">
          <div><strong id="nl-title"></strong><br><small class="text-muted">Priority: <span id="nl-priority"></span> · Deadline: <span id="nl-deadline"></span></small></div>
          <span id="nl-confidence" class="badge"></span>
        </div>
        <a id="nl-create-link" href="#" class="btn btn-purple btn-sm mt-2">Create Task</a>
      </div>
    </div>
  </div>
</div>

<!-- Filters -->
<form class="row g-2 mb-3" method="GET">
  <div class="col-auto">
    <select name="status" class="form-select form-select-sm">
      <option value="">All Statuses</option>
      {% for s in ['Pending','In Progress','Completed'] %}
      <option value="{{ s }}" {{ 'selected' if request.args.get('status') == s else '' }}>{{ s }}</option>
      {% endfor %}
    </select>
  </div>
  <div class="col-auto">
    <select name="priority" class="form-select form-select-sm">
      <option value="">All Priorities</option>
      {% for p in ['Low','Medium','High','Critical'] %}
      <option value="{{ p }}" {{ 'selected' if request.args.get('priority') == p else '' }}>{{ p }}</option>
      {% endfor %}
    </select>
  </div>
  <div class="col-auto"><button class="btn btn-sm btn-outline-secondary" type="submit">Filter</button></div>
</form>

<div class="table-responsive">
  <table class="table table-hover align-middle">
    <thead class="table-light">
      <tr><th>Title</th><th>Priority</th><th>Status</th><th>Deadline</th><th>Time</th><th></th></tr>
    </thead>
    <tbody>
      {% for t in tasks %}
      <tr>
        <td>{{ t.title }}</td>
        <td><span class="priority-{{ t.priority }}">{{ t.priority }}</span></td>
        <td>
          <span class="badge bg-{{ 'success' if t.status == 'Completed' else ('warning text-dark' if t.status == 'In Progress' else 'secondary') }}">
            {{ t.status }}
          </span>
        </td>
        <td>{{ t.deadline.strftime('%b %d, %Y %H:%M') if t.deadline else '—' }}</td>
        <td><small>{{ t.time_logged_minutes or 0 }}m</small></td>
        <td class="text-end">
          <a href="{{ url_for('task_subtasks', task_id=t.id) }}" class="btn btn-sm btn-outline-info" title="Subtasks"><i class="bi bi-list-check"></i></a>
          <a href="{{ url_for('tasks_edit', task_id=t.id) }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-pencil"></i></a>
          <form method="POST" action="{{ url_for('tasks_delete', task_id=t.id) }}" class="d-inline"
                onsubmit="return confirm('Delete this task?')">
            {{ csrf_token_field|safe }}
            <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
          </form>
        </td>
      </tr>
      {% else %}
      <tr><td colspan="6" class="text-center text-muted">No tasks found.</td></tr>
      {% endfor %}
    </tbody>
  </table>
</div>
{% endblock %}
""").replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script>
let nlTimer;
const nlInput = document.getElementById('nl-input');
if(nlInput){
  nlInput.addEventListener('input', ()=>{
    clearTimeout(nlTimer);
    nlTimer = setTimeout(()=>{
      const text = nlInput.value.trim();
      if(!text){ document.getElementById('nl-preview').style.display='none'; return; }
      fetch("{{ url_for('tasks_parse_nl') }}",{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({text})})
      .then(r=>r.json()).then(d=>{
        document.getElementById('nl-preview').style.display='block';
        document.getElementById('nl-title').textContent = d.title;
        document.getElementById('nl-priority').textContent = d.priority;
        document.getElementById('nl-deadline').textContent = d.deadline ? new Date(d.deadline).toLocaleDateString() : 'Not detected';
        const conf = document.getElementById('nl-confidence');
        conf.textContent = Math.round(d.confidence*100)+'% confidence';
        conf.className = 'badge bg-'+(d.confidence>=0.8?'success':d.confidence>=0.5?'warning text-dark':'danger');
        let url = "{{ url_for('tasks_new') }}?title="+encodeURIComponent(d.title)+"&priority="+encodeURIComponent(d.priority);
        if(d.deadline) url += "&deadline="+encodeURIComponent(d.deadline.slice(0,16));
        document.getElementById('nl-create-link').href = url;
      });
    }, 400);
  });
}
</script>
{% endblock %}""")

TASK_FORM_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-7">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0">{{ 'Edit' if task else 'New' }} Task</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.title.label(class="form-label") }}
            {{ form.title(class="form-control" + (" is-invalid" if form.title.errors else "")) }}
            {% for e in form.title.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.description.label(class="form-label") }}
            {{ form.description(class="form-control", rows=3) }}
          </div>
          <div class="row g-3">
            <div class="col-md-6">
              {{ form.priority.label(class="form-label") }}
              {{ form.priority(class="form-select") }}
            </div>
            <div class="col-md-6">
              {{ form.status.label(class="form-label") }}
              {{ form.status(class="form-select") }}
            </div>
          </div>
          <div class="mb-3 mt-3">
            {{ form.deadline.label(class="form-label") }}
            {{ form.deadline(class="form-control" + (" is-invalid" if form.deadline.errors else "")) }}
            {% for e in form.deadline.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.recurrence.label(class="form-label") }}
            {{ form.recurrence(class="form-select") }}
          </div>
          <div class="d-flex gap-2">
            {{ form.submit(class="btn btn-primary") }}
            <a href="{{ url_for('tasks_index') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
        {% if task %}
        <hr>
        <div class="card card-dark"><div class="card-body py-2">
          <label class="form-label small">Log Time</label>
          <form method="POST" action="{{ url_for('task_log_time', task_id=task.id) }}" class="row g-2 align-items-end">
            {{ csrf_token_field|safe }}
            <div class="col-auto"><input name="minutes" type="number" min="1" class="form-control form-control-sm" placeholder="Minutes" required></div>
            <div class="col"><input name="note" class="form-control form-control-sm" placeholder="Note (optional)"></div>
            <div class="col-auto"><button class="btn btn-sm btn-teal">Log</button></div>
          </form>
          <small class="text-muted">Total logged: {{ task.time_logged_minutes or 0 }} minutes · <a href="{{ url_for('task_subtasks', task_id=task.id) }}">Manage subtasks</a></small>
        </div></div>
        {% endif %}
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── Goals templates ──────────────────────────────────────────────────────────

GOALS_INDEX_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-trophy"></i> Goals</h3>
  <a href="{{ url_for('goals_new') }}" class="btn btn-primary"><i class="bi bi-plus-lg"></i> New Goal</a>
</div>
<div class="row g-3">
  {% for g in goals %}
  <div class="col-md-6 col-lg-4">
    <div class="card h-100 border-{{ 'success' if g.status == 'Achieved' else ('danger' if g.status == 'Overdue' else 'primary') }}">
      <div class="card-body">
        <div class="d-flex justify-content-between">
          <h6 class="card-title">{{ g.name }}</h6>
          <span class="badge bg-{{ 'success' if g.status == 'Achieved' else ('danger' if g.status == 'Overdue' else 'primary') }}">{{ g.status }}</span>
        </div>
        <p class="card-text text-muted small">{{ g.description[:80] + '...' if g.description and g.description|length > 80 else g.description or '' }}</p>
        <div class="mb-2">
          <div class="d-flex justify-content-between"><small>Progress</small><small>{{ g.progress }}%</small></div>
          <div class="progress" style="height:8px">
            <div class="progress-bar bg-{{ 'success' if g.status == 'Achieved' else ('danger' if g.status == 'Overdue' else 'primary') }}"
                 style="width:{{ g.progress }}%"></div>
          </div>
        </div>
        <small class="text-muted">Target: {{ g.target_date.strftime('%b %d, %Y') }}</small>
      </div>
      <div class="card-footer bg-transparent d-flex gap-2">
        <a href="{{ url_for('goals_edit', goal_id=g.id) }}" class="btn btn-sm btn-outline-secondary">Edit</a>
        <form method="POST" action="{{ url_for('goals_delete', goal_id=g.id) }}" class="d-inline"
              onsubmit="return confirm('Delete this goal?')">
          {{ csrf_token_field|safe }}
          <button class="btn btn-sm btn-outline-danger">Delete</button>
        </form>
      </div>
    </div>
  </div>
  {% else %}
  <div class="col"><p class="text-muted">No goals yet. <a href="{{ url_for('goals_new') }}">Create one!</a></p></div>
  {% endfor %}
</div>
{% endblock %}
""")

GOAL_FORM_TEMPLATE = (
    BASE_TEMPLATE
    .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="row justify-content-center">
  <div class="col-md-7">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0">{{ 'Edit' if goal else 'New' }} Goal</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.name.label(class="form-label") }}
            {{ form.name(class="form-control" + (" is-invalid" if form.name.errors else "")) }}
            {% for e in form.name.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.description.label(class="form-label") }}
            {{ form.description(class="form-control", rows=3) }}
          </div>
          <div class="row g-3 mb-3">
            <div class="col-md-6">
              {{ form.target_date.label(class="form-label") }}
              {{ form.target_date(class="form-control" + (" is-invalid" if form.target_date.errors else "")) }}
              {% for e in form.target_date.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
            </div>
            <div class="col-md-6">
              {{ form.progress.label(class="form-label") }}
              {{ form.progress(class="form-control", type="range", min=0, max=100) }}
              <div class="form-text" id="prog-val">{{ form.progress.data or 0 }}%</div>
            </div>
          </div>
          <div class="d-flex gap-2">
            {{ form.submit(class="btn btn-primary") }}
            <a href="{{ url_for('goals_index') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}""")
    .replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script>
const slider = document.getElementById('progress');
const label = document.getElementById('prog-val');
if(slider && label){ slider.addEventListener('input', ()=>{ label.textContent = slider.value+'%'; }); }
</script>
{% endblock %}""")
)

# ── Calendar template ────────────────────────────────────────────────────────

CALENDAR_INDEX_TEMPLATE = (
    BASE_TEMPLATE
    .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-calendar3"></i> Calendar</h3>
  <a href="{{ url_for('calendar_new') }}" class="btn btn-primary"><i class="bi bi-plus-lg"></i> New Event</a>
</div>
<div id="calendar"></div>
<hr>
<h5>All Events</h5>
<div class="table-responsive">
  <table class="table table-hover align-middle">
    <thead class="table-light"><tr><th>Title</th><th>Type</th><th>Start</th><th>End</th><th></th></tr></thead>
    <tbody>
      {% for e in events %}
      <tr>
        <td>{{ e.title }}</td>
        <td><span class="badge bg-info text-dark">{{ e.event_type }}</span></td>
        <td>{{ e.start_datetime.strftime('%b %d, %Y %H:%M') }}</td>
        <td>{{ e.end_datetime.strftime('%b %d, %Y %H:%M') if e.end_datetime else '—' }}</td>
        <td class="text-end">
          <a href="{{ url_for('calendar_edit', event_id=e.id) }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-pencil"></i></a>
          <form method="POST" action="{{ url_for('calendar_delete', event_id=e.id) }}" class="d-inline"
                onsubmit="return confirm('Delete event?')">
            {{ csrf_token_field|safe }}
            <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
          </form>
        </td>
      </tr>
      {% else %}
      <tr><td colspan="5" class="text-center text-muted">No events yet.</td></tr>
      {% endfor %}
    </tbody>
  </table>
</div>
{% endblock %}""")
    .replace("{% block extra_css %}{% endblock %}", """{% block extra_css %}
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/fullcalendar@6.1.11/index.global.min.css">
{% endblock %}""")
    .replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script src="https://cdn.jsdelivr.net/npm/fullcalendar@6.1.11/index.global.min.js"></script>
<script>
document.addEventListener('DOMContentLoaded', function(){
  const cal = new FullCalendar.Calendar(document.getElementById('calendar'), {
    initialView: 'dayGridMonth',
    height: 500,
    events: {{ events_json|safe }},
  });
  cal.render();
});
</script>
{% endblock %}""")
)

CALENDAR_FORM_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-7">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0">{{ 'Edit' if event else 'New' }} Event</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.title.label(class="form-label") }}
            {{ form.title(class="form-control" + (" is-invalid" if form.title.errors else "")) }}
            {% for e in form.title.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.description.label(class="form-label") }}
            {{ form.description(class="form-control", rows=2) }}
          </div>
          <div class="row g-3 mb-3">
            <div class="col-md-6">
              {{ form.start_datetime.label(class="form-label") }}
              {{ form.start_datetime(class="form-control" + (" is-invalid" if form.start_datetime.errors else "")) }}
              {% for e in form.start_datetime.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
            </div>
            <div class="col-md-6">
              {{ form.end_datetime.label(class="form-label") }}
              {{ form.end_datetime(class="form-control") }}
            </div>
          </div>
          <div class="mb-3">
            {{ form.event_type.label(class="form-label") }}
            {{ form.event_type(class="form-select") }}
          </div>
          <div class="d-flex gap-2">
            {{ form.submit(class="btn btn-primary") }}
            <a href="{{ url_for('calendar_index') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── Notes templates ──────────────────────────────────────────────────────────

NOTES_INDEX_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-journal-text"></i> Notes</h3>
  <div class="d-flex gap-2">
    <a href="{{ url_for('notes_graph') }}" class="btn btn-outline-secondary btn-sm"><i class="bi bi-diagram-3"></i> Knowledge Graph</a>
    <a href="{{ url_for('notes_new') }}" class="btn btn-primary"><i class="bi bi-plus-lg"></i> New Note</a>
  </div>
</div>
<div class="row g-3">
  {% for n in notes %}
  <div class="col-md-6 col-lg-4">
    <div class="card h-100">
      <div class="card-body">
        <h6 class="card-title">{{ n.title }}</h6>
        {% if n.category %}<span class="badge bg-secondary mb-2">{{ n.category }}</span>{% endif %}
        <p class="card-text text-muted small">{{ n.content[:120] + '...' if n.content and n.content|length > 120 else n.content or '' }}</p>
        <small class="text-muted">{{ n.updated_at.strftime('%b %d, %Y') }}</small>
      </div>
      <div class="card-footer bg-transparent d-flex gap-2">
        <a href="{{ url_for('notes_edit', note_id=n.id) }}" class="btn btn-sm btn-outline-secondary">Edit</a>
        <form method="POST" action="{{ url_for('notes_delete', note_id=n.id) }}" class="d-inline"
              onsubmit="return confirm('Delete this note?')">
          {{ csrf_token_field|safe }}
          <button class="btn btn-sm btn-outline-danger">Delete</button>
        </form>
      </div>
    </div>
  </div>
  {% else %}
  <div class="col"><p class="text-muted">No notes yet. <a href="{{ url_for('notes_new') }}">Create one!</a></p></div>
  {% endfor %}
</div>
{% endblock %}
""")

NOTE_FORM_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-8">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0">{{ 'Edit' if note else 'New' }} Note</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.title.label(class="form-label") }}
            {{ form.title(class="form-control" + (" is-invalid" if form.title.errors else "")) }}
            {% for e in form.title.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.category.label(class="form-label") }}
            {{ form.category(class="form-control") }}
          </div>
          <div class="mb-3">
            {{ form.content.label(class="form-label") }}
            {{ form.content(class="form-control", rows=10) }}
          </div>
          <div class="d-flex gap-2">
            {{ form.submit(class="btn btn-primary") }}
            <a href="{{ url_for('notes_index') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── Documents templates ──────────────────────────────────────────────────────

DOCUMENTS_INDEX_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-folder2"></i> Document Vault</h3>
  <a href="{{ url_for('documents_upload') }}" class="btn btn-primary"><i class="bi bi-upload"></i> Upload</a>
</div>
<div class="table-responsive">
  <table class="table table-hover align-middle">
    <thead class="table-light"><tr><th>Name</th><th>Type</th><th>Size</th><th>Uploaded</th><th>Extracted</th><th></th></tr></thead>
    <tbody>
      {% for d in documents %}
      <tr>
        <td>{{ d.original_name }}</td>
        <td><span class="badge bg-secondary">{{ d.file_type.upper() }}</span></td>
        <td>{{ '%.1f'|format(d.file_size / 1024) }} KB</td>
        <td>{{ d.upload_date.strftime('%b %d, %Y') }}</td>
        <td>
          {% if d.parse_failed %}<span class="text-danger small">Parse failed</span>
          {% elif d.extracted_text %}<span class="text-success small">✓ Text available</span>
          {% else %}<span class="text-muted small">—</span>{% endif %}
        </td>
        <td class="text-end">
          <a href="{{ url_for('documents_view', doc_id=d.id) }}" class="btn btn-sm btn-outline-info"><i class="bi bi-eye"></i></a>
          <a href="{{ url_for('documents_download', doc_id=d.id) }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-download"></i></a>
          <form method="POST" action="{{ url_for('documents_delete', doc_id=d.id) }}" class="d-inline"
                onsubmit="return confirm('Delete this document?')">
            {{ csrf_token_field|safe }}
            <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
          </form>
        </td>
      </tr>
      {% else %}
      <tr><td colspan="6" class="text-center text-muted">No documents uploaded yet.</td></tr>
      {% endfor %}
    </tbody>
  </table>
</div>
{% endblock %}
""")

DOCUMENTS_UPLOAD_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-6">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0"><i class="bi bi-upload"></i> Upload Document</h5></div>
      <div class="card-body">
        <form method="POST" enctype="multipart/form-data" novalidate>
          {{ csrf_form.hidden_tag() }}
          <div class="mb-3">
            <label class="form-label">File <span class="text-muted small">(PDF, DOCX, JPG, PNG — max 10 MB)</span></label>
            <input type="file" name="file" class="form-control" required accept=".pdf,.docx,.jpg,.jpeg,.png">
          </div>
          <div class="d-flex gap-2">
            <button type="submit" class="btn btn-primary">Upload</button>
            <a href="{{ url_for('documents_index') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

DOCUMENT_VIEW_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-lg-8">
    <div class="card shadow-sm">
      <div class="card-header d-flex justify-content-between align-items-center">
        <h5 class="mb-0"><i class="bi bi-file-text"></i> {{ doc.original_name }}</h5>
        <a href="{{ url_for('documents_download', doc_id=doc.id) }}" class="btn btn-sm btn-outline-primary">Download</a>
      </div>
      <div class="card-body">
        <dl class="row">
          <dt class="col-sm-3">File Type</dt><dd class="col-sm-9">{{ doc.file_type.upper() }}</dd>
          <dt class="col-sm-3">Size</dt><dd class="col-sm-9">{{ '%.2f'|format(doc.file_size / 1024) }} KB</dd>
          <dt class="col-sm-3">Uploaded</dt><dd class="col-sm-9">{{ doc.upload_date.strftime('%B %d, %Y %H:%M') }}</dd>
        </dl>
        {% if doc.parse_failed %}
          <div class="alert alert-warning">Text extraction failed for this file.</div>
        {% elif doc.extracted_text %}
          <h6>Extracted Text</h6>
          <pre class="bg-light p-3 rounded" style="max-height:400px;overflow-y:auto;white-space:pre-wrap">{{ doc.extracted_text[:5000] }}</pre>
          {% if doc.extracted_text|length > 5000 %}<p class="text-muted small">Showing first 5000 characters.</p>{% endif %}
        {% else %}
          <p class="text-muted">No text extracted (image or unsupported type).</p>
        {% endif %}
      </div>
      <div class="card-footer"><a href="{{ url_for('documents_index') }}" class="btn btn-outline-secondary">Back</a></div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── Reminders template ───────────────────────────────────────────────────────

REMINDERS_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <h3><i class="bi bi-bell"></i> Reminders</h3>
  <a href="{{ url_for('reminders_new') }}" class="btn btn-primary"><i class="bi bi-plus-lg"></i> New Reminder</a>
</div>
<div class="row g-3">
  <div class="col-lg-8">
    <div class="card">
      <div class="card-header">All Reminders</div>
      <div class="list-group list-group-flush">
        {% for r in reminders %}
        <div class="list-group-item d-flex justify-content-between align-items-start {{ 'list-group-item-light' if r.is_read else '' }}">
          <div>
            <div class="fw-{{ 'normal' if r.is_read else 'semibold' }}">{{ r.text }}</div>
            <small class="text-muted"><i class="bi bi-clock"></i> {{ r.remind_at.strftime('%b %d, %Y %H:%M') }}</small>
            {% if not r.is_read and r.remind_at <= now %}
              <span class="badge bg-danger ms-2">Due</span>
            {% endif %}
          </div>
          <div class="d-flex gap-1">
            {% if not r.is_read %}
            <form method="POST" action="{{ url_for('reminders_dismiss', reminder_id=r.id) }}">
              {{ csrf_token_field|safe }}
              <button class="btn btn-sm btn-outline-success" title="Mark read"><i class="bi bi-check-lg"></i></button>
            </form>
            {% endif %}
            <form method="POST" action="{{ url_for('reminders_delete', reminder_id=r.id) }}"
                  onsubmit="return confirm('Delete?')">
              {{ csrf_token_field|safe }}
              <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
            </form>
          </div>
        </div>
        {% else %}
        <div class="list-group-item text-muted">No reminders set.</div>
        {% endfor %}
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

REMINDER_FORM_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="row justify-content-center">
  <div class="col-md-6">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0">New Reminder</h5></div>
      <div class="card-body">
        <form method="POST" novalidate>
          {{ form.hidden_tag() }}
          <div class="mb-3">
            {{ form.text.label(class="form-label") }}
            {{ form.text(class="form-control" + (" is-invalid" if form.text.errors else "")) }}
            {% for e in form.text.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="mb-3">
            {{ form.remind_at.label(class="form-label") }}
            {{ form.remind_at(class="form-control" + (" is-invalid" if form.remind_at.errors else "")) }}
            {% for e in form.remind_at.errors %}<div class="invalid-feedback">{{ e }}</div>{% endfor %}
          </div>
          <div class="d-flex gap-2">
            {{ form.submit(class="btn btn-primary") }}
            <a href="{{ url_for('reminders_panel') }}" class="btn btn-outline-secondary">Cancel</a>
          </div>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
""")

# ── AI Assistant template ────────────────────────────────────────────────────

AI_CHAT_TEMPLATE = (
    BASE_TEMPLATE
    .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="row justify-content-center">
  <div class="col-lg-8">
    <div class="card shadow-sm">
      <div class="card-header"><h5 class="mb-0"><i class="bi bi-robot"></i> AI Assistant</h5></div>
      <div class="card-body" style="min-height:350px">
        <div id="chat-log" class="mb-3" style="max-height:400px;overflow-y:auto">
          {% for msg in history %}
          <div class="mb-2 {{ 'd-flex justify-content-end' if msg.role == 'user' else '' }}">
            <div class="p-2 rounded"
                 style="max-width:80%;white-space:pre-wrap;{{ 'background:var(--accent-purple);color:white' if msg.role == 'user' else 'background:#1f2937;color:var(--text-primary);border:1px solid var(--border-color)' }}">{{ msg.text }}</div>
          </div>
          {% endfor %}
        </div>
        <form method="POST" novalidate>
          {{ csrf_form.hidden_tag() }}
          <div class="input-group">
            <input type="text" name="query" id="query-input" class="form-control" placeholder="Ask me anything about your tasks, goals, or schedule…"
                   value="{{ last_query|default('') }}" maxlength="500" required autocomplete="off">
            <button type="button" id="voice-btn" class="btn btn-outline-secondary" title="Voice input">
              <i class="bi bi-mic" id="mic-icon"></i>
            </button>
            <button class="btn btn-primary" type="submit"><i class="bi bi-send"></i> Ask</button>
          </div>
          <div id="voice-status" style="font-size:0.75rem;color:var(--text-secondary);margin-top:4px;display:none">
            🎙️ Listening... speak now
          </div>
        </form>
        {% if not ai_available %}
        <div class="alert alert-warning mt-3 mb-0">
          <i class="bi bi-exclamation-triangle"></i>
          AI model not available. Install <code>sentence-transformers</code> to enable semantic matching.
          Falling back to keyword matching.
        </div>
        {% endif %}
      </div>
    </div>

    {% if recommendations %}
    <div class="card mt-3">
      <div class="card-header"><i class="bi bi-lightbulb"></i> Recommended Focus</div>
      <ul class="list-group list-group-flush">
        {% for t in recommendations %}
        <li class="list-group-item d-flex justify-content-between">
          <span>{{ t.title }}</span>
          <span class="badge bg-secondary">{{ t.priority }}</span>
        </li>
        {% endfor %}
      </ul>
    </div>
    {% endif %}
  </div>
</div>
{% endblock %}""")
    .replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script>
// auto-scroll chat log
const log = document.getElementById('chat-log');
if(log) log.scrollTop = log.scrollHeight;

// Voice input using Web Speech API
const voiceBtn = document.getElementById('voice-btn');
const micIcon = document.getElementById('mic-icon');
const voiceStatus = document.getElementById('voice-status');
const queryInput = document.getElementById('query-input');

if (voiceBtn) {
  const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
  if (SpeechRecognition) {
    const recognition = new SpeechRecognition();
    recognition.lang = 'en-US';
    recognition.interimResults = false;
    recognition.maxAlternatives = 1;

    let listening = false;

    voiceBtn.addEventListener('click', () => {
      if (listening) {
        recognition.stop();
        return;
      }
      recognition.start();
    });

    recognition.onstart = () => {
      listening = true;
      micIcon.className = 'bi bi-mic-fill';
      voiceBtn.classList.replace('btn-outline-secondary', 'btn-danger');
      if (voiceStatus) { voiceStatus.style.display = 'block'; }
    };

    recognition.onresult = (event) => {
      const transcript = event.results[0][0].transcript;
      if (queryInput) queryInput.value = transcript;
    };

    recognition.onend = () => {
      listening = false;
      micIcon.className = 'bi bi-mic';
      voiceBtn.classList.replace('btn-danger', 'btn-outline-secondary');
      if (voiceStatus) { voiceStatus.style.display = 'none'; }
    };

    recognition.onerror = (event) => {
      listening = false;
      micIcon.className = 'bi bi-mic';
      voiceBtn.classList.replace('btn-danger', 'btn-outline-secondary');
      if (voiceStatus) { voiceStatus.style.display = 'none'; }
      if (event.error !== 'no-speech') {
        alert('Voice error: ' + event.error + '. Try Chrome or Edge browser.');
      }
    };
  } else {
    voiceBtn.title = 'Voice input not supported in this browser (use Chrome or Edge)';
    voiceBtn.disabled = true;
  }
}
</script>
{% endblock %}""")
)

# ── Analytics template ───────────────────────────────────────────────────────

ANALYTICS_TEMPLATE = (
    BASE_TEMPLATE
    .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="d-flex align-items-center gap-2 mb-4">
  <i class="bi bi-bar-chart-line" style="font-size:1.4rem;color:var(--accent-purple)"></i>
  <h3 class="mb-0">Analytics</h3>
</div>

<!-- Summary cards -->
<div class="row g-3 mb-4">
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card purple">
      <span class="stat-num" style="color:#a78bfa">{{ stats.total_tasks }}</span>
      <span class="stat-label">Total Tasks</span>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card green">
      <span class="stat-num" style="color:#34d399">{{ stats.completed_tasks }}</span>
      <span class="stat-label">Completed</span>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card amber">
      <span class="stat-num" style="color:#fbbf24">{{ '%.0f'|format(stats.completion_rate) }}%</span>
      <span class="stat-label">Completion Rate</span>
    </div>
  </div>
  <div class="col-sm-6 col-lg-3">
    <div class="stat-card red">
      <span class="stat-num" style="color:#f87171">{{ '%.0f'|format(stats.productivity_score) }}</span>
      <span class="stat-label">Productivity Score</span>
    </div>
  </div>
</div>

<div class="row g-4">
  <!-- Weekly completions chart -->
  <div class="col-lg-7">
    <div class="card">
      <div class="card-header"><i class="bi bi-bar-chart me-1" style="color:var(--accent-purple)"></i>Weekly Task Completions (last 8 weeks)</div>
      <div class="card-body"><canvas id="weeklyChart" style="max-height:200px"></canvas></div>
    </div>
  </div>

  <!-- AI Reports donut: task status distribution -->
  <div class="col-lg-5">
    <div class="card">
      <div class="card-header"><i class="bi bi-pie-chart me-1" style="color:var(--accent-pink)"></i>AI Report: Task Status</div>
      <div class="card-body"><canvas id="taskStatusChart" height="180"></canvas></div>
    </div>
  </div>

  <!-- Goal status doughnut -->
  <div class="col-lg-4">
    <div class="card">
      <div class="card-header"><i class="bi bi-trophy me-1" style="color:var(--accent-green)"></i>Goal Status</div>
      <div class="card-body"><canvas id="goalChart" style="max-height:220px"></canvas></div>
    </div>
  </div>

  <!-- Priority breakdown -->
  <div class="col-lg-4">
    <div class="card">
      <div class="card-header"><i class="bi bi-funnel me-1" style="color:var(--accent-amber)"></i>Tasks by Priority</div>
      <div class="card-body"><canvas id="priorityChart" style="max-height:200px"></canvas></div>
    </div>
  </div>

  <!-- Habit completion rate + goal ratios -->
  <div class="col-lg-4">
    <div class="card">
      <div class="card-header"><i class="bi bi-activity me-1" style="color:var(--accent-teal)"></i>Habit & Goal Health</div>
      <div class="card-body">
        <div class="mb-3">
          <div class="d-flex justify-content-between mb-1"><span style="font-size:0.85rem">Habit Completion (30d)</span><strong style="color:var(--accent-teal)">{{ stats.habit_rate }}%</strong></div>
          <div class="progress" style="height:6px"><div class="progress-bar" style="width:{{ stats.habit_rate }}%;background:var(--accent-teal)"></div></div>
        </div>
        <div class="mb-3">
          <div class="d-flex justify-content-between mb-1"><span style="font-size:0.85rem">Goals Achieved</span><strong style="color:var(--accent-green)">{{ stats.achieved_goals }}</strong></div>
          <div class="progress" style="height:6px"><div class="progress-bar" style="width:{{ stats.goal_achievement_ratio }}%;background:var(--accent-green)"></div></div>
        </div>
        <div class="mb-3">
          <div class="d-flex justify-content-between mb-1"><span style="font-size:0.85rem">Goals Active</span><strong style="color:#a78bfa">{{ stats.active_goals }}</strong></div>
          <div class="progress" style="height:6px"><div class="progress-bar" style="width:{{ (stats.active_goals / (stats.total_goals or 1) * 100)|round }}%;background:var(--accent-purple)"></div></div>
        </div>
        <div>
          <div class="d-flex justify-content-between mb-1"><span style="font-size:0.85rem">Goals Overdue</span><strong style="color:var(--accent-red)">{{ stats.overdue_goals }}</strong></div>
          <div class="progress" style="height:6px"><div class="progress-bar" style="width:{{ (stats.overdue_goals / (stats.total_goals or 1) * 100)|round }}%;background:var(--accent-red)"></div></div>
        </div>
        <hr style="border-color:var(--border-color);margin:12px 0">
        <p class="mb-0" style="font-size:0.85rem">Avg goal progress: <strong style="color:var(--accent-amber)">{{ '%.0f'|format(stats.avg_goal_progress) }}%</strong></p>
      </div>
    </div>
  </div>
</div>
{% endblock %}""")
    .replace("{% block extra_js %}{% endblock %}", """{% block extra_js %}
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.3/dist/chart.umd.min.js"></script>
<script>
const chartDefaults = {
  plugins: { legend: { labels: { color: '#8b949e', font: { size: 11 } } } },
  scales: {
    x: { grid: { color: '#30363d' }, ticks: { color: '#8b949e' } },
    y: { grid: { color: '#30363d' }, ticks: { color: '#8b949e', stepSize: 1 }, beginAtZero: true }
  }
};

const wData = {{ weekly_data|safe }};
new Chart(document.getElementById('weeklyChart'), {
  type: 'bar',
  data: {
    labels: wData.labels,
    datasets: [{ label: 'Completed Tasks', data: wData.values, backgroundColor: '#7c3aed', borderRadius: 5, borderSkipped: false }]
  },
  options: { responsive: true, maintainAspectRatio: false, scales: { y: { beginAtZero: true, ticks: { stepSize: 1, color: '#8b949e' }, grid: { color: '#30363d' } }, x: { ticks: { color: '#8b949e' }, grid: { color: '#30363d' } } }, plugins: { legend: { labels: { color: '#8b949e' } } } }
});

const tsData = {{ task_status_data|safe }};
new Chart(document.getElementById('taskStatusChart'), {
  type: 'doughnut',
  data: {
    labels: ['Pending', 'In Progress', 'Completed'],
    datasets: [{ data: tsData, backgroundColor: ['#7c3aed', '#f59e0b', '#10b981'], borderWidth: 0, hoverOffset: 8 }]
  },
  options: { plugins: { legend: { position: 'bottom', labels: { color: '#8b949e' } } }, cutout: '65%' }
});

const gData = {{ goal_data|safe }};
new Chart(document.getElementById('goalChart'), {
  type: 'doughnut',
  data: {
    labels: ['Achieved', 'Active', 'Overdue'],
    datasets: [{ data: gData, backgroundColor: ['#10b981', '#7c3aed', '#ef4444'], borderWidth: 0, hoverOffset: 8 }]
  },
  options: { responsive: true, maintainAspectRatio: false, plugins: { legend: { position: 'bottom', labels: { color: '#8b949e' } } } }
});

const pData = {{ priority_data|safe }};
new Chart(document.getElementById('priorityChart'), {
  type: 'bar',
  data: {
    labels: ['Critical', 'High', 'Medium', 'Low'],
    datasets: [{ label: 'Tasks', data: pData, backgroundColor: ['#ef4444', '#f59e0b', '#7c3aed', '#8b949e'], borderRadius: 5, borderSkipped: false }]
  },
  options: { responsive: true, maintainAspectRatio: false, indexAxis: 'y', scales: { x: { beginAtZero: true, ticks: { stepSize: 1, color: '#8b949e' }, grid: { color: '#30363d' } }, y: { ticks: { color: '#8b949e' }, grid: { color: '#30363d' } } }, plugins: { legend: { display: false } } }
});
</script>
{% endblock %}""")
)

# ── Projects template ────────────────────────────────────────────────────────

PROJECTS_INDEX_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """{% block content %}
<div class="d-flex justify-content-between align-items-center mb-4">
  <div>
    <h3 class="mb-0"><i class="bi bi-briefcase" style="color:var(--accent-purple)"></i> Projects & Portfolio</h3>
    <p class="text-muted small mb-0">Your resume-ready project entries</p>
  </div>
  <a href="{{ url_for('projects_new') }}" class="btn btn-purple btn-sm"><i class="bi bi-plus-lg"></i> Add Project</a>
</div>

{% if projects %}
<div class="row g-4">
  {% for p in projects %}
  <div class="col-12">
    <div class="card-dark p-4 rounded-3">
      <div class="d-flex justify-content-between align-items-start flex-wrap gap-2 mb-2">
        <div>
          <h5 class="mb-0" style="color:var(--text-primary)">{{ p.title }}</h5>
          <div class="d-flex align-items-center gap-2 mt-1 flex-wrap">
            <span class="badge" style="background:#7c3aed33;color:#a78bfa;border:1px solid #7c3aed55;font-size:0.72rem">{{ p.project_type }}</span>
            {% if p.start_date %}
            <span style="font-size:0.8rem;color:var(--text-secondary)">
              <i class="bi bi-calendar3"></i>
              {{ p.start_date }}{% if p.end_date %} &ndash; {{ p.end_date }}{% endif %}
            </span>
            {% endif %}
          </div>
        </div>
        <div class="d-flex gap-2 align-items-center">
          {% if p.github_url %}
          <a href="{{ p.github_url }}" target="_blank" class="btn btn-sm btn-outline-secondary" title="GitHub">
            <i class="bi bi-github"></i> GitHub
          </a>
          {% endif %}
          {% if p.live_url %}
          <a href="{{ p.live_url }}" target="_blank" class="btn btn-sm btn-outline-secondary" title="Live Demo">
            <i class="bi bi-box-arrow-up-right"></i> Live
          </a>
          {% endif %}
          <a href="{{ url_for('projects_edit', project_id=p.id) }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-pencil"></i></a>
          <form method="POST" action="{{ url_for('projects_delete', project_id=p.id) }}" class="d-inline"
                onsubmit="return confirm('Delete this project?')">
            {{ csrf_field|safe }}
            <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
          </form>
        </div>
      </div>

      {% if p.description %}
      <ul class="mb-2" style="padding-left:1.2rem">
        {% for line in p.description.strip().split('\\n') %}
        {% if line.strip() %}
        <li style="font-size:0.88rem;color:var(--text-primary);margin-bottom:3px">{{ line.strip().lstrip('•-') }}</li>
        {% endif %}
        {% endfor %}
      </ul>
      {% endif %}

      {% if p.technologies %}
      <div class="d-flex flex-wrap gap-1 mt-2">
        {% for tech in p.technologies.split(',') %}
        <span style="background:#06b6d422;color:#67e8f9;border:1px solid #06b6d444;border-radius:6px;padding:2px 8px;font-size:0.72rem">{{ tech.strip() }}</span>
        {% endfor %}
      </div>
      {% endif %}

      {% if p.achievement %}
      <div class="mt-2" style="font-size:0.82rem;color:var(--accent-amber)">
        <i class="bi bi-trophy-fill"></i> {{ p.achievement }}
      </div>
      {% endif %}
    </div>
  </div>
  {% endfor %}
</div>
{% else %}
<div class="text-center mt-5" style="color:var(--text-secondary)">
  <i class="bi bi-briefcase" style="font-size:3rem;opacity:0.3"></i>
  <p class="mt-3">No projects yet. <a href="{{ url_for('projects_new') }}">Add your first project!</a></p>
</div>
{% endif %}
{% endblock %}""")

PROJECT_FORM_TEMPLATE = BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """{% block content %}
<div class="row justify-content-center">
  <div class="col-lg-7">
    <div class="card p-4">
      <h5 class="mb-4"><i class="bi bi-briefcase" style="color:var(--accent-purple)"></i> {{ 'Edit' if project else 'New' }} Project</h5>
      <form method="POST" novalidate>
        {{ form.hidden_tag() }}
        <div class="row g-3 mb-3">
          <div class="col-md-8">
            {{ form.title.label(class="form-label") }}
            {{ form.title(class="form-control") }}
            {% for e in form.title.errors %}<div class="text-danger small">{{ e }}</div>{% endfor %}
          </div>
          <div class="col-md-4">
            {{ form.project_type.label(class="form-label") }}
            {{ form.project_type(class="form-select") }}
          </div>
        </div>
        <div class="row g-3 mb-3">
          <div class="col-md-6">
            {{ form.start_date.label(class="form-label") }}
            {{ form.start_date(class="form-control", placeholder="e.g. Aug 2025") }}
          </div>
          <div class="col-md-6">
            {{ form.end_date.label(class="form-label") }}
            {{ form.end_date(class="form-control", placeholder="e.g. Present or Dec 2025") }}
          </div>
        </div>
        <div class="row g-3 mb-3">
          <div class="col-md-6">
            {{ form.github_url.label(class="form-label") }}
            {{ form.github_url(class="form-control", placeholder="https://github.com/...") }}
          </div>
          <div class="col-md-6">
            {{ form.live_url.label(class="form-label") }}
            {{ form.live_url(class="form-control", placeholder="https://...") }}
          </div>
        </div>
        <div class="mb-3">
          {{ form.description.label(class="form-label") }}
          {{ form.description(class="form-control", rows=5, placeholder="• Developed AI-powered resume builder\n• Integrated OpenAI API for suggestions\n• Achieved 2nd Prize at HackFest 2025") }}
          <div class="form-text">One bullet point per line. Start each with • or just write the text.</div>
        </div>
        <div class="mb-3">
          {{ form.technologies.label(class="form-label") }}
          {{ form.technologies(class="form-control", placeholder="React.js, Node.js, TypeScript, Tailwind CSS, OpenAI API") }}
          <div class="form-text">Comma-separated list of technologies</div>
        </div>
        <div class="mb-4">
          {{ form.achievement.label(class="form-label") }}
          {{ form.achievement(class="form-control", placeholder="2nd Prize at HackFest Hackathon, SGBTI College, 2025") }}
        </div>
        <div class="d-flex gap-2">
          {{ form.submit(class="btn btn-purple") }}
          <a href="{{ url_for('projects_index') }}" class="btn btn-outline-secondary">Cancel</a>
        </div>
      </form>
    </div>
  </div>
</div>
{% endblock %}""")


def _csrf_hidden() -> str:
    """Return a hidden CSRF input tag as a string (used in template contexts)."""
    from flask_wtf.csrf import generate_csrf
    token = generate_csrf()
    return f'<input type="hidden" name="csrf_token" value="{token}">'


class _CsrfForm(FlaskForm):
    """Minimal form just for the hidden CSRF tag."""
    pass


# =============================================================================
# Route registration
# =============================================================================

def _register_routes(app: Flask):

    # ──────────────────────────────────────────────────────────────────────────
    # Auth
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/register", methods=["GET", "POST"])
    def register():
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        form = RegistrationForm()
        if form.validate_on_submit():
            user = User(name=_sanitize(form.name.data), email=form.email.data.lower())
            user.set_password(form.password.data)
            db.session.add(user)
            db.session.commit()
            login_user(user)
            flash("Account created. Welcome!", "success")
            return redirect(url_for("dashboard"))
        return render_template_string(REGISTER_TEMPLATE, form=form)

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        form = LoginForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data.lower()).first()
            if user and user.check_password(form.password.data):
                login_user(user, remember=form.remember.data)
                next_page = request.args.get("next")
                flash("Logged in.", "success")
                return redirect(next_page or url_for("dashboard"))
            flash("Invalid email or password.", "danger")
        return render_template_string(LOGIN_TEMPLATE, form=form)

    @app.route("/logout")
    @login_required
    def logout():
        logout_user()
        flash("You have been logged out.", "info")
        return redirect(url_for("login"))

    def _generate_reset_token(user):
        raw = secrets.token_urlsafe(32)
        token_hash = hashlib.sha256(raw.encode()).hexdigest()
        expires_at = datetime.utcnow() + timedelta(minutes=60)
        PasswordResetToken.query.filter_by(user_id=user.id, used=False).update({"used": True})
        token_record = PasswordResetToken(user_id=user.id, token_hash=token_hash, expires_at=expires_at)
        db.session.add(token_record)
        db.session.commit()
        return raw

    @app.route("/reset-request", methods=["GET", "POST"])
    def reset_request():
        if current_user.is_authenticated:
            return redirect(url_for("dashboard"))
        form = PasswordResetRequestForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data.lower()).first()
            if user:
                raw_token = _generate_reset_token(user)
                reset_url = url_for("reset_password", token=raw_token, _external=True)
                try:
                    msg = Message("Password Reset Request", recipients=[user.email])
                    msg.body = (
                        f"Click the link to reset your password:\n{reset_url}\n"
                        "This link expires in 60 minutes."
                    )
                    mail.send(msg)
                except Exception:
                    pass
            flash("If that email is registered, a reset link has been sent.", "info")
            return redirect(url_for("login"))
        return render_template_string(RESET_REQUEST_TEMPLATE, form=form)

    @app.route("/reset-password/<token>", methods=["GET", "POST"])
    def reset_password(token):
        token_hash = hashlib.sha256(token.encode()).hexdigest()
        token_record = PasswordResetToken.query.filter_by(token_hash=token_hash, used=False).first()
        if not token_record or token_record.expires_at < datetime.utcnow():
            flash("The reset link is invalid or has expired.", "danger")
            return redirect(url_for("reset_request"))
        form = PasswordResetForm()
        if form.validate_on_submit():
            user = User.query.get(token_record.user_id)
            user.set_password(form.password.data)
            token_record.used = True
            db.session.commit()
            flash("Password updated. Please log in.", "success")
            return redirect(url_for("login"))
        return render_template_string(RESET_PASSWORD_TEMPLATE, form=form)

    @app.route("/profile", methods=["GET", "POST"])
    @login_required
    def profile():
        form = ProfileForm(obj=current_user)
        if form.validate_on_submit():
            new_email = form.email.data.lower()
            if new_email != current_user.email:
                if User.query.filter_by(email=new_email).first():
                    flash("Email already in use.", "danger")
                    return render_template_string(PROFILE_TEMPLATE, form=form)
            current_user.name = _sanitize(form.name.data)
            current_user.email = new_email
            current_user.theme = form.theme.data
            current_user.email_digest = form.email_digest.data
            db.session.commit()
            flash("Profile updated.", "success")
            return redirect(url_for("profile"))
        return render_template_string(PROFILE_TEMPLATE, form=form)

    @app.route("/delete-account", methods=["POST"])
    @login_required
    def delete_account():
        user = current_user._get_current_object()
        logout_user()
        db.session.delete(user)
        db.session.commit()
        flash("Account deleted.", "info")
        return redirect(url_for("register"))

    # ──────────────────────────────────────────────────────────────────────────
    # Dashboard
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/")
    @login_required
    def dashboard():
        uid = current_user.id
        now = datetime.utcnow()
        today = date.today()
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)

        # Stats
        total_tasks = Task.query.filter_by(user_id=uid).count()
        pending = Task.query.filter_by(user_id=uid, status="Pending").count()
        completed_tasks = Task.query.filter_by(user_id=uid, status="Completed").count()
        due_today = Task.query.filter(
            Task.user_id == uid,
            Task.deadline >= today_start,
            Task.deadline < today_end,
            Task.status != "Completed",
        ).count()
        overdue = Task.query.filter(
            Task.user_id == uid,
            Task.deadline < now,
            Task.status != "Completed",
        ).count()
        goals_active = Goal.query.filter_by(user_id=uid, status="Active").count()
        total_goals = Goal.query.filter_by(user_id=uid).count()
        achieved_goals = Goal.query.filter_by(user_id=uid, status="Achieved").count()

        stats = {
            "total_tasks": total_tasks,
            "pending": pending,
            "due_today": due_today,
            "overdue": overdue,
            "goals_active": goals_active,
        }

        # AI Score
        ai_score = int(
            (completed_tasks / max(total_tasks, 1)) * 60
            + (achieved_goals / max(total_goals, 1)) * 40
        )

        # Greeting — use local time, not UTC
        local_hour = datetime.now().hour
        if local_hour < 5:
            greeting = "Good night 🌙"
        elif local_hour < 12:
            greeting = "Good morning ☀️"
        elif local_hour < 17:
            greeting = "Good afternoon 🌤️"
        elif local_hour < 21:
            greeting = "Good evening 🌆"
        else:
            greeting = "Good night 🌙"

        # Upcoming deadlines (tasks with deadline in next 7 days)
        upcoming_deadlines = Task.query.filter(
            Task.user_id == uid,
            Task.deadline.isnot(None),
            Task.deadline >= now,
            Task.deadline <= now + timedelta(days=7),
            Task.status != "Completed",
        ).order_by(Task.deadline).limit(5).all()

        # Goals for progress display
        goals = Goal.query.filter_by(user_id=uid).order_by(Goal.target_date).limit(4).all()

        # Recent notes
        recent_notes = Note.query.filter_by(user_id=uid).order_by(Note.updated_at.desc()).limit(5).all()

        # AI recommendations
        recommendations = _recommend_tasks(uid)

        # Habits
        habits = Habit.query.filter_by(user_id=uid).all()
        checked_today = {
            c.habit_id for c in HabitCheckin.query.filter_by(
                user_id=uid, checked_date=today
            ).all()
        }

        # Weekly activity chart (last 7 days completed tasks)
        weekly_labels = []
        weekly_values = []
        for i in range(6, -1, -1):
            d = today - timedelta(days=i)
            d_start = datetime.combine(d, datetime.min.time())
            d_end = d_start + timedelta(days=1)
            count = Task.query.filter(
                Task.user_id == uid,
                Task.status == "Completed",
                Task.completed_at >= d_start,
                Task.completed_at < d_end,
            ).count()
            weekly_labels.append(d.strftime("%a"))
            weekly_values.append(count)

        return render_template_string(
            DASHBOARD_TEMPLATE,
            stats=stats,
            ai_score=ai_score,
            greeting=greeting,
            upcoming_deadlines=upcoming_deadlines,
            goals=goals,
            recent_notes=recent_notes,
            recommendations=recommendations,
            habits=habits,
            checked_today=checked_today,
            weekly_labels=json.dumps(weekly_labels),
            weekly_values=json.dumps(weekly_values),
        )

    # ──────────────────────────────────────────────────────────────────────────
    # Tasks
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/tasks")
    @login_required
    def tasks_index():
        uid = current_user.id
        q = Task.query.filter_by(user_id=uid)
        status_filter = request.args.get("status")
        priority_filter = request.args.get("priority")
        if status_filter:
            q = q.filter_by(status=status_filter)
        if priority_filter:
            q = q.filter_by(priority=priority_filter)
        tasks = q.order_by(Task.created_at.desc()).all()
        return render_template_string(
            TASKS_INDEX_TEMPLATE, tasks=tasks,
            csrf_token_field=_csrf_hidden(),
        )

    @app.route("/tasks/new", methods=["GET", "POST"])
    @login_required
    def tasks_new():
        form = TaskForm()
        if request.args.get("title"):
            form.title.data = request.args.get("title")
        if request.args.get("priority"):
            form.priority.data = request.args.get("priority")
        if request.args.get("deadline"):
            try:
                form.deadline.data = datetime.fromisoformat(request.args.get("deadline"))
            except ValueError:
                pass
        if form.validate_on_submit():
            t = Task(
                user_id=current_user.id,
                title=_sanitize(form.title.data),
                description=_sanitize(form.description.data or ""),
                priority=form.priority.data,
                status=form.status.data,
                deadline=form.deadline.data,
                recurrence=form.recurrence.data,
            )
            if t.status == "Completed" and not t.completed_at:
                t.completed_at = datetime.utcnow()
            db.session.add(t)
            db.session.commit()
            flash("Task created.", "success")
            return redirect(url_for("tasks_index"))
        return render_template_string(TASK_FORM_TEMPLATE, form=form, task=None)

    @app.route("/tasks/<int:task_id>/edit", methods=["GET", "POST"])
    @login_required
    def tasks_edit(task_id):
        t = Task.query.get_or_404(task_id)
        if t.user_id != current_user.id:
            abort(403)
        form = TaskForm(obj=t)
        if form.validate_on_submit():
            t.title = _sanitize(form.title.data)
            t.description = _sanitize(form.description.data or "")
            t.priority = form.priority.data
            old_status = t.status
            t.status = form.status.data
            t.deadline = form.deadline.data
            t.recurrence = form.recurrence.data
            if t.status == "Completed" and old_status != "Completed":
                t.completed_at = datetime.utcnow()
            elif t.status != "Completed":
                t.completed_at = None
            db.session.commit()
            flash("Task updated.", "success")
            return redirect(url_for("tasks_index"))
        return render_template_string(TASK_FORM_TEMPLATE, form=form, task=t)

    @app.route("/tasks/<int:task_id>/delete", methods=["POST"])
    @login_required
    def tasks_delete(task_id):
        t = Task.query.get_or_404(task_id)
        if t.user_id != current_user.id:
            abort(403)
        db.session.delete(t)
        db.session.commit()
        flash("Task deleted.", "info")
        return redirect(url_for("tasks_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Goals
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/goals")
    @login_required
    def goals_index():
        goals = Goal.query.filter_by(user_id=current_user.id).order_by(Goal.target_date).all()
        return render_template_string(GOALS_INDEX_TEMPLATE, goals=goals, csrf_token_field=_csrf_hidden())

    @app.route("/goals/new", methods=["GET", "POST"])
    @login_required
    def goals_new():
        form = GoalForm()
        if form.validate_on_submit():
            g = Goal(
                user_id=current_user.id,
                name=_sanitize(form.name.data),
                description=_sanitize(form.description.data or ""),
                target_date=form.target_date.data,
            )
            g.update_progress(form.progress.data or 0)
            db.session.add(g)
            db.session.commit()
            flash("Goal created.", "success")
            return redirect(url_for("goals_index"))
        return render_template_string(GOAL_FORM_TEMPLATE, form=form, goal=None)

    @app.route("/goals/<int:goal_id>/edit", methods=["GET", "POST"])
    @login_required
    def goals_edit(goal_id):
        g = Goal.query.get_or_404(goal_id)
        if g.user_id != current_user.id:
            abort(403)
        form = GoalForm(obj=g)
        if form.validate_on_submit():
            g.name = _sanitize(form.name.data)
            g.description = _sanitize(form.description.data or "")
            g.target_date = form.target_date.data
            g.update_progress(form.progress.data or 0)
            db.session.commit()
            flash("Goal updated.", "success")
            return redirect(url_for("goals_index"))
        return render_template_string(GOAL_FORM_TEMPLATE, form=form, goal=g)

    @app.route("/goals/<int:goal_id>/delete", methods=["POST"])
    @login_required
    def goals_delete(goal_id):
        g = Goal.query.get_or_404(goal_id)
        if g.user_id != current_user.id:
            abort(403)
        db.session.delete(g)
        db.session.commit()
        flash("Goal deleted.", "info")
        return redirect(url_for("goals_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Calendar
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/calendar")
    @login_required
    def calendar_index():
        events = (CalendarEvent.query
                  .filter_by(user_id=current_user.id)
                  .order_by(CalendarEvent.start_datetime)
                  .all())
        events_json = json.dumps([
            {
                "title": e.title,
                "start": e.start_datetime.isoformat(),
                "end": e.end_datetime.isoformat() if e.end_datetime else None,
                "color": "#0d6efd" if e.event_type == "Event" else ("#dc3545" if e.event_type == "TaskDeadline" else "#198754"),
            }
            for e in events
        ])
        return render_template_string(
            CALENDAR_INDEX_TEMPLATE, events=events,
            events_json=events_json, csrf_token_field=_csrf_hidden(),
        )

    @app.route("/calendar/new", methods=["GET", "POST"])
    @login_required
    def calendar_new():
        form = CalendarEventForm()
        if form.validate_on_submit():
            e = CalendarEvent(
                user_id=current_user.id,
                title=_sanitize(form.title.data),
                description=_sanitize(form.description.data or ""),
                start_datetime=form.start_datetime.data,
                end_datetime=form.end_datetime.data,
                event_type=form.event_type.data,
            )
            db.session.add(e)
            db.session.commit()
            flash("Event created.", "success")
            return redirect(url_for("calendar_index"))
        return render_template_string(CALENDAR_FORM_TEMPLATE, form=form, event=None)

    @app.route("/calendar/<int:event_id>/edit", methods=["GET", "POST"])
    @login_required
    def calendar_edit(event_id):
        e = CalendarEvent.query.get_or_404(event_id)
        if e.user_id != current_user.id:
            abort(403)
        form = CalendarEventForm(obj=e)
        if form.validate_on_submit():
            e.title = _sanitize(form.title.data)
            e.description = _sanitize(form.description.data or "")
            e.start_datetime = form.start_datetime.data
            e.end_datetime = form.end_datetime.data
            e.event_type = form.event_type.data
            db.session.commit()
            flash("Event updated.", "success")
            return redirect(url_for("calendar_index"))
        return render_template_string(CALENDAR_FORM_TEMPLATE, form=form, event=e)

    @app.route("/calendar/<int:event_id>/delete", methods=["POST"])
    @login_required
    def calendar_delete(event_id):
        e = CalendarEvent.query.get_or_404(event_id)
        if e.user_id != current_user.id:
            abort(403)
        db.session.delete(e)
        db.session.commit()
        flash("Event deleted.", "info")
        return redirect(url_for("calendar_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Notes
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/notes")
    @login_required
    def notes_index():
        notes = (Note.query.filter_by(user_id=current_user.id)
                 .order_by(Note.updated_at.desc()).all())
        return render_template_string(NOTES_INDEX_TEMPLATE, notes=notes, csrf_token_field=_csrf_hidden())

    @app.route("/notes/new", methods=["GET", "POST"])
    @login_required
    def notes_new():
        form = NoteForm()
        if form.validate_on_submit():
            n = Note(
                user_id=current_user.id,
                title=_sanitize(form.title.data),
                content=_sanitize(form.content.data or ""),
                category=_sanitize(form.category.data or ""),
            )
            db.session.add(n)
            db.session.commit()
            if _v2_ctx.get("recompute_note_links"):
                threading.Thread(
                    target=_v2_ctx["recompute_note_links"],
                    args=(current_user.id, n.id, app),
                    daemon=True,
                ).start()
            flash("Note saved.", "success")
            return redirect(url_for("notes_index"))
        return render_template_string(NOTE_FORM_TEMPLATE, form=form, note=None)

    @app.route("/notes/<int:note_id>/edit", methods=["GET", "POST"])
    @login_required
    def notes_edit(note_id):
        n = Note.query.get_or_404(note_id)
        if n.user_id != current_user.id:
            abort(403)
        form = NoteForm(obj=n)
        if form.validate_on_submit():
            n.title = _sanitize(form.title.data)
            n.content = _sanitize(form.content.data or "")
            n.category = _sanitize(form.category.data or "")
            n.updated_at = datetime.utcnow()
            db.session.commit()
            if _v2_ctx.get("recompute_note_links"):
                threading.Thread(
                    target=_v2_ctx["recompute_note_links"],
                    args=(current_user.id, n.id, app),
                    daemon=True,
                ).start()
            flash("Note updated.", "success")
            return redirect(url_for("notes_index"))
        return render_template_string(NOTE_FORM_TEMPLATE, form=form, note=n)

    @app.route("/notes/<int:note_id>/delete", methods=["POST"])
    @login_required
    def notes_delete(note_id):
        n = Note.query.get_or_404(note_id)
        if n.user_id != current_user.id:
            abort(403)
        db.session.delete(n)
        db.session.commit()
        flash("Note deleted.", "info")
        return redirect(url_for("notes_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Documents
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/documents")
    @login_required
    def documents_index():
        docs = (Document.query.filter_by(user_id=current_user.id)
                .order_by(Document.upload_date.desc()).all())
        return render_template_string(
            DOCUMENTS_INDEX_TEMPLATE, documents=docs, csrf_token_field=_csrf_hidden()
        )

    @app.route("/documents/upload", methods=["GET", "POST"])
    @login_required
    def documents_upload():
        csrf_form = _CsrfForm()
        if request.method == "POST" and csrf_form.validate_on_submit():
            file = request.files.get("file")
            if not file or file.filename == "":
                flash("No file selected.", "danger")
                return redirect(request.url)
            if not _allowed_file(file.filename, app.config["ALLOWED_EXTENSIONS"]):
                flash("File type not allowed. Use PDF, DOCX, JPG, or PNG.", "danger")
                return redirect(request.url)
            original_name = secure_filename(file.filename)
            ext = original_name.rsplit(".", 1)[1].lower()
            stored_name = f"{uuid.uuid4().hex}_{original_name}"
            folder = _upload_dir(current_user.id)
            filepath = os.path.join(folder, stored_name)
            file.save(filepath)
            file_size = os.path.getsize(filepath)
            extracted_text, parse_failed = _extract_text(filepath, ext)
            doc = Document(
                user_id=current_user.id,
                filename=stored_name,
                original_name=original_name,
                file_size=file_size,
                file_type=ext,
                extracted_text=extracted_text,
                parse_failed=parse_failed,
            )
            db.session.add(doc)
            db.session.commit()
            flash("Document uploaded successfully.", "success")
            return redirect(url_for("documents_index"))
        return render_template_string(DOCUMENTS_UPLOAD_TEMPLATE, csrf_form=csrf_form)

    @app.route("/documents/<int:doc_id>")
    @login_required
    def documents_view(doc_id):
        doc = Document.query.get_or_404(doc_id)
        if doc.user_id != current_user.id:
            abort(403)
        return render_template_string(DOCUMENT_VIEW_TEMPLATE, doc=doc)

    @app.route("/documents/<int:doc_id>/download")
    @login_required
    def documents_download(doc_id):
        doc = Document.query.get_or_404(doc_id)
        if doc.user_id != current_user.id:
            abort(403)
        folder = _upload_dir(current_user.id)
        return send_from_directory(folder, doc.filename, as_attachment=True,
                                   download_name=doc.original_name)

    @app.route("/documents/<int:doc_id>/delete", methods=["POST"])
    @login_required
    def documents_delete(doc_id):
        doc = Document.query.get_or_404(doc_id)
        if doc.user_id != current_user.id:
            abort(403)
        # Remove file from disk
        folder = _upload_dir(current_user.id)
        filepath = os.path.join(folder, doc.filename)
        try:
            os.remove(filepath)
        except OSError:
            pass
        db.session.delete(doc)
        db.session.commit()
        flash("Document deleted.", "info")
        return redirect(url_for("documents_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Reminders
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/reminders")
    @login_required
    def reminders_panel():
        reminders = (Reminder.query.filter_by(user_id=current_user.id)
                     .order_by(Reminder.remind_at).all())
        return render_template_string(
            REMINDERS_TEMPLATE,
            reminders=reminders,
            now=datetime.utcnow(),
            csrf_token_field=_csrf_hidden(),
        )

    @app.route("/reminders/new", methods=["GET", "POST"])
    @login_required
    def reminders_new():
        form = ReminderForm()
        if form.validate_on_submit():
            r = Reminder(
                user_id=current_user.id,
                text=_sanitize(form.text.data),
                remind_at=form.remind_at.data,
                source_type="Manual",
            )
            db.session.add(r)
            db.session.commit()
            flash("Reminder set.", "success")
            return redirect(url_for("reminders_panel"))
        return render_template_string(REMINDER_FORM_TEMPLATE, form=form)

    @app.route("/reminders/<int:reminder_id>/dismiss", methods=["POST"])
    @login_required
    def reminders_dismiss(reminder_id):
        r = Reminder.query.get_or_404(reminder_id)
        if r.user_id != current_user.id:
            abort(403)
        r.is_read = True
        db.session.commit()
        return redirect(url_for("reminders_panel"))

    @app.route("/reminders/<int:reminder_id>/delete", methods=["POST"])
    @login_required
    def reminders_delete(reminder_id):
        r = Reminder.query.get_or_404(reminder_id)
        if r.user_id != current_user.id:
            abort(403)
        db.session.delete(r)
        db.session.commit()
        flash("Reminder deleted.", "info")
        return redirect(url_for("reminders_panel"))

    @app.route("/reminders/unread-count")
    @login_required
    def reminders_unread_count():
        count = Reminder.query.filter_by(
            user_id=current_user.id,
            is_read=False,
        ).filter(Reminder.remind_at <= datetime.utcnow()).count()
        return jsonify({"count": count})

    # ──────────────────────────────────────────────────────────────────────────
    # AI Assistant
    # ──────────────────────────────────────────────────────────────────────────

    _chat_history: dict = {}  # user_id -> list of {role, text}

    @app.route("/ai", methods=["GET", "POST"])
    @login_required
    def ai_chat():
        uid = current_user.id
        csrf_form = _CsrfForm()
        history = _chat_history.setdefault(uid, [])
        last_query = ""
        recommendations = _recommend_tasks(uid)

        if request.method == "POST" and csrf_form.validate_on_submit():
            query = request.form.get("query", "").strip()[:500]
            last_query = query
            if query:
                history.append({"role": "user", "text": query})
                # Try semantic intent matching first
                intent = _match_intent(query)
                if intent:
                    answer = _handle_intent(intent, uid)
                else:
                    # Keyword fallback
                    ql = query.lower()
                    if any(w in ql for w in ["pending", "todo", "to do", "not done", "work", "remaining", "left", "incomplete", "unfinished", "what do i have", "my tasks", "open task"]):
                        answer = _handle_intent("pending_tasks", uid)
                    elif any(w in ql for w in ["deadline", "due", "overdue", "expir", "upcoming", "soon", "this week", "next", "schedule"]):
                        answer = _handle_intent("upcoming_deadlines", uid)
                    elif any(w in ql for w in ["goal", "progress", "achieve", "target", "milestone", "objective", "on track"]):
                        answer = _handle_intent("goal_progress", uid)
                    elif any(w in ql for w in ["focus", "today", "priority", "important", "urgent", "first", "start", "begin", "concentrate"]):
                        answer = _handle_intent("focus_today", uid)
                    elif any(w in ql for w in ["remind", "reminder", "notification", "alert"]):
                        reminders = Reminder.query.filter_by(user_id=uid, is_read=False).filter(
                            Reminder.remind_at <= datetime.utcnow()
                        ).order_by(Reminder.remind_at.desc()).limit(5).all()
                        if reminders:
                            lines = [f"• {r.text}" for r in reminders]
                            answer = "Your active reminders:\n" + "\n".join(lines)
                        else:
                            answer = "You have no active reminders right now."
                    elif any(w in ql for w in ["habit", "streak", "checkin", "check in", "daily"]):
                        habits = Habit.query.filter_by(user_id=uid).order_by(Habit.streak.desc()).limit(5).all()
                        if habits:
                            lines = [f"• {h.name} — 🔥 {h.streak} day streak" for h in habits]
                            answer = "Your habits:\n" + "\n".join(lines)
                        else:
                            answer = "You haven't created any habits yet. Go to Habits to start tracking!"
                    elif any(w in ql for w in ["note", "notes"]):
                        notes = Note.query.filter_by(user_id=uid).order_by(Note.updated_at.desc()).limit(5).all()
                        if notes:
                            lines = [f"• {n.title}" + (f" [{n.category}]" if n.category else "") for n in notes]
                            answer = "Your recent notes:\n" + "\n".join(lines)
                        else:
                            answer = "You have no notes yet."
                    elif any(w in ql for w in ["hello", "hi", "hey", "help", "what can you do", "how are you"]):
                        answer = (
                            "Hi! I'm your AI assistant. Here's what I can help with:\n"
                            "• 'What are my pending tasks?' — show open tasks\n"
                            "• 'Show upcoming deadlines' — tasks due soon\n"
                            "• 'How are my goals going?' — goal progress\n"
                            "• 'What should I focus on today?' — top priorities\n"
                            "• 'Show my reminders' — active notifications\n"
                            "• 'My habits' — habit streaks\n"
                            "• 'My notes' — recent notes\n"
                            "\nYou can also use the 🎙️ mic button to speak your question!"
                        )
                    else:
                        answer = (
                            "I didn't quite understand that. Try asking:\n"
                            "• 'What are my pending tasks?'\n"
                            "• 'Show upcoming deadlines'\n"
                            "• 'How are my goals going?'\n"
                            "• 'What should I focus on today?'\n"
                            "• 'Show my reminders'\n"
                            "Or click the 🎙️ mic to speak your question!"
                        )
                history.append({"role": "assistant", "text": answer})
                # Keep last 20 exchanges
                _chat_history[uid] = history[-40:]

        return render_template_string(
            AI_CHAT_TEMPLATE,
            history=history,
            last_query=last_query,
            recommendations=recommendations,
            ai_available=_AI_AVAILABLE,
            csrf_form=csrf_form,
        )

    # ──────────────────────────────────────────────────────────────────────────
    # Analytics
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/analytics")
    @login_required
    def analytics_index():
        uid = current_user.id
        now = datetime.utcnow()

        # ── Task statistics ────────────────────────────────────────────────
        total_tasks = Task.query.filter_by(user_id=uid).count()
        completed_tasks = Task.query.filter_by(user_id=uid, status="Completed").count()
        completion_rate = (completed_tasks / total_tasks * 100) if total_tasks else 0

        # Priority breakdown [Critical, High, Medium, Low]
        priority_data = [
            Task.query.filter_by(user_id=uid, priority="Critical").count(),
            Task.query.filter_by(user_id=uid, priority="High").count(),
            Task.query.filter_by(user_id=uid, priority="Medium").count(),
            Task.query.filter_by(user_id=uid, priority="Low").count(),
        ]

        # ── Weekly completions (last 8 weeks) ──────────────────────────────
        weekly_labels = []
        weekly_values = []
        for i in range(7, -1, -1):
            week_start = (now - timedelta(weeks=i)).replace(
                hour=0, minute=0, second=0, microsecond=0
            )
            week_end = week_start + timedelta(weeks=1)
            count = Task.query.filter(
                Task.user_id == uid,
                Task.completed_at >= week_start,
                Task.completed_at < week_end,
            ).count()
            weekly_labels.append(week_start.strftime("W%W"))
            weekly_values.append(count)
        weekly_data = json.dumps({"labels": weekly_labels, "values": weekly_values})

        # ── Goal statistics ────────────────────────────────────────────────
        total_goals = Goal.query.filter_by(user_id=uid).count()
        achieved_goals = Goal.query.filter_by(user_id=uid, status="Achieved").count()
        active_goals = Goal.query.filter_by(user_id=uid, status="Active").count()
        overdue_goals = Goal.query.filter_by(user_id=uid, status="Overdue").count()
        goal_achievement_ratio = (achieved_goals / total_goals * 100) if total_goals else 0
        goal_data = json.dumps([achieved_goals, active_goals, overdue_goals])

        all_goals = Goal.query.filter_by(user_id=uid).all()
        avg_goal_progress = (
            sum(g.progress for g in all_goals) / len(all_goals) if all_goals else 0
        )

        # ── Productivity score (heuristic) ─────────────────────────────────
        # Score = completion_rate * 0.5 + goal_achievement_ratio * 0.3 + (avg_goal_progress / 100) * 20
        productivity_score = (
            completion_rate * 0.5
            + goal_achievement_ratio * 0.3
            + (avg_goal_progress / 100) * 20
        )

        # ── Task status breakdown [Pending, In Progress, Completed] ────────
        pending_tasks = Task.query.filter_by(user_id=uid, status="Pending").count()
        inprogress_tasks = Task.query.filter_by(user_id=uid, status="In Progress").count()
        task_status_data = json.dumps([pending_tasks, inprogress_tasks, completed_tasks])

        # ── Habit completion rate (last 30 days) ────────────────────────────
        thirty_days_ago = (now - timedelta(days=30)).date()
        all_habits = Habit.query.filter_by(user_id=uid).all()
        if all_habits:
            total_possible = len(all_habits) * 30
            total_checkins = HabitCheckin.query.filter(
                HabitCheckin.user_id == uid,
                HabitCheckin.checked_date >= thirty_days_ago,
            ).count()
            habit_rate = round(total_checkins / total_possible * 100) if total_possible else 0
        else:
            habit_rate = 0

        stats = {
            "total_tasks": total_tasks,
            "completed_tasks": completed_tasks,
            "completion_rate": completion_rate,
            "total_goals": total_goals,
            "achieved_goals": achieved_goals,
            "active_goals": active_goals,
            "overdue_goals": overdue_goals,
            "goal_achievement_ratio": goal_achievement_ratio,
            "avg_goal_progress": avg_goal_progress,
            "productivity_score": min(productivity_score, 100),
            "habit_rate": habit_rate,
        }

        return render_template_string(
            ANALYTICS_TEMPLATE,
            stats=stats,
            weekly_data=weekly_data,
            goal_data=goal_data,
            priority_data=json.dumps(priority_data),
            task_status_data=task_status_data,
        )

    # ──────────────────────────────────────────────────────────────────────────
    # Habits
    # ──────────────────────────────────────────────────────────────────────────

    HABITS_INDEX_TEMPLATE = (
        BASE_TEMPLATE
        .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="d-flex justify-content-between align-items-center mb-4">
  <h3><i class="bi bi-activity"></i> Habit Tracker</h3>
  <a href="{{ url_for('habits_new') }}" class="btn btn-primary btn-sm"><i class="bi bi-plus-lg"></i> New Habit</a>
</div>

{% if habits %}
<div class="row g-3 mb-4">
  {% for h in habits %}
  <div class="col-md-6 col-lg-4">
    <div class="card-dark p-3 rounded-3">
      <div class="d-flex justify-content-between align-items-start mb-2">
        <div>
          <span class="fw-bold" style="color:{{ h.color }}">{{ h.name }}</span>
          <div class="text-secondary small">{{ h.frequency }}</div>
        </div>
        <div class="text-center">
          <div class="fs-4">🔥</div>
          <div class="fw-bold" style="color:#f59e0b">{{ h.streak }}</div>
          <div class="text-secondary" style="font-size:0.7rem">streak</div>
        </div>
      </div>
      <div class="d-flex gap-1 mb-3">
        {% for day in h.week_status %}
        <div style="width:28px;height:28px;border-radius:6px;background:{{ '#10b981' if day else '#30363d' }};display:flex;align-items:center;justify-content:center;font-size:0.7rem;color:white">
          {{ loop.index0 | weekday_letter }}
        </div>
        {% endfor %}
      </div>
      <div class="d-flex gap-2">
        {% if h.id in checked_today %}
        <span class="badge bg-success px-3 py-2">✓ Done Today</span>
        {% else %}
        <form method="POST" action="{{ url_for('habits_checkin', habit_id=h.id) }}" class="d-inline">
          {{ csrf_field|safe }}
          <button class="btn btn-sm btn-outline-success">Check In</button>
        </form>
        {% endif %}
        <a href="{{ url_for('habits_edit', habit_id=h.id) }}" class="btn btn-sm btn-outline-secondary"><i class="bi bi-pencil"></i></a>
        <form method="POST" action="{{ url_for('habits_delete', habit_id=h.id) }}" class="d-inline"
              onsubmit="return confirm('Delete habit?')">
          {{ csrf_field|safe }}
          <button class="btn btn-sm btn-outline-danger"><i class="bi bi-trash"></i></button>
        </form>
      </div>
    </div>
  </div>
  {% endfor %}
</div>
{% else %}
<div class="text-center text-secondary mt-5">
  <i class="bi bi-activity fs-1"></i>
  <p class="mt-2">No habits yet. <a href="{{ url_for('habits_new') }}">Create your first habit!</a></p>
</div>
{% endif %}
{% endblock %}""")
    )

    HABIT_FORM_TEMPLATE = (
        BASE_TEMPLATE
        .replace("{% block content %}{% endblock %}", """{% block content %}
<div class="row justify-content-center">
  <div class="col-md-6">
    <div class="card-dark p-4 rounded-3">
      <h5 class="mb-4">{{ 'Edit' if habit else 'New' }} Habit</h5>
      <form method="POST" novalidate>
        {{ form.hidden_tag() }}
        <div class="mb-3">
          {{ form.name.label(class="form-label") }}
          {{ form.name(class="form-control") }}
          {% for e in form.name.errors %}<div class="text-danger small">{{ e }}</div>{% endfor %}
        </div>
        <div class="mb-3">
          {{ form.description.label(class="form-label") }}
          {{ form.description(class="form-control", rows=3) }}
        </div>
        <div class="mb-3">
          {{ form.frequency.label(class="form-label") }}
          {{ form.frequency(class="form-select") }}
        </div>
        <div class="mb-3">
          {{ form.color.label(class="form-label") }}
          {{ form.color(class="form-control form-control-color", type="color") }}
        </div>
        <div class="d-flex gap-2">
          {{ form.submit(class="btn btn-primary") }}
          <a href="{{ url_for('habits_index') }}" class="btn btn-outline-secondary">Cancel</a>
        </div>
      </form>
    </div>
  </div>
</div>
{% endblock %}""")
    )

    @app.route("/habits")
    @login_required
    def habits_index():
        from datetime import timedelta
        today = date.today()
        habits = Habit.query.filter_by(user_id=current_user.id).all()
        checked_today = {
            c.habit_id for c in HabitCheckin.query.filter_by(
                user_id=current_user.id, checked_date=today
            ).all()
        }
        # Build 7-day status for each habit
        for h in habits:
            week_status = []
            for i in range(6, -1, -1):
                d = today - timedelta(days=i)
                checked = HabitCheckin.query.filter_by(
                    habit_id=h.id, checked_date=d
                ).first() is not None
                week_status.append(checked)
            h.week_status = week_status
        return render_template_string(
            HABITS_INDEX_TEMPLATE,
            habits=habits,
            checked_today=checked_today,
            csrf_field=_csrf_hidden(),
        )

    @app.route("/habits/new", methods=["GET", "POST"])
    @login_required
    def habits_new():
        form = HabitForm()
        if form.validate_on_submit():
            h = Habit(
                user_id=current_user.id,
                name=_sanitize(form.name.data),
                description=_sanitize(form.description.data or ""),
                frequency=form.frequency.data,
                color=form.color.data or "#7c3aed",
            )
            db.session.add(h)
            db.session.commit()
            flash("Habit created.", "success")
            return redirect(url_for("habits_index"))
        return render_template_string(HABIT_FORM_TEMPLATE, form=form, habit=None)

    @app.route("/habits/<int:habit_id>/edit", methods=["GET", "POST"])
    @login_required
    def habits_edit(habit_id):
        h = Habit.query.get_or_404(habit_id)
        if h.user_id != current_user.id:
            abort(403)
        form = HabitForm(obj=h)
        if form.validate_on_submit():
            h.name = _sanitize(form.name.data)
            h.description = _sanitize(form.description.data or "")
            h.frequency = form.frequency.data
            h.color = form.color.data or "#7c3aed"
            db.session.commit()
            flash("Habit updated.", "success")
            return redirect(url_for("habits_index"))
        return render_template_string(HABIT_FORM_TEMPLATE, form=form, habit=h)

    @app.route("/habits/<int:habit_id>/delete", methods=["POST"])
    @login_required
    def habits_delete(habit_id):
        h = Habit.query.get_or_404(habit_id)
        if h.user_id != current_user.id:
            abort(403)
        db.session.delete(h)
        db.session.commit()
        flash("Habit deleted.", "info")
        return redirect(url_for("habits_index"))

    @app.route("/habits/<int:habit_id>/checkin", methods=["POST"])
    @login_required
    def habits_checkin(habit_id):
        h = Habit.query.get_or_404(habit_id)
        if h.user_id != current_user.id:
            abort(403)
        today = date.today()
        existing = HabitCheckin.query.filter_by(
            habit_id=habit_id, checked_date=today
        ).first()
        if not existing:
            checkin = HabitCheckin(
                habit_id=habit_id,
                user_id=current_user.id,
                checked_date=today,
            )
            db.session.add(checkin)
            # Update streak
            yesterday = today - timedelta(days=1)
            prev = HabitCheckin.query.filter_by(
                habit_id=habit_id, checked_date=yesterday
            ).first()
            h.streak = (h.streak or 0) + 1 if prev else 1
            h.last_checked = today
            db.session.commit()
            flash(f"✓ '{h.name}' checked in! Streak: {h.streak} days 🔥", "success")
        else:
            flash("Already checked in today.", "info")
        return redirect(url_for("habits_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Projects
    # ──────────────────────────────────────────────────────────────────────────

    @app.route("/projects")
    @login_required
    def projects_index():
        projects = Project.query.filter_by(user_id=current_user.id).order_by(Project.created_at.desc()).all()
        return render_template_string(PROJECTS_INDEX_TEMPLATE, projects=projects, csrf_field=_csrf_hidden())

    @app.route("/projects/new", methods=["GET", "POST"])
    @login_required
    def projects_new():
        form = ProjectForm()
        if form.validate_on_submit():
            p = Project(
                user_id=current_user.id,
                title=_sanitize(form.title.data),
                project_type=form.project_type.data,
                start_date=_sanitize(form.start_date.data or ""),
                end_date=_sanitize(form.end_date.data or ""),
                github_url=_sanitize(form.github_url.data or ""),
                live_url=_sanitize(form.live_url.data or ""),
                description=_sanitize(form.description.data or ""),
                technologies=_sanitize(form.technologies.data or ""),
                achievement=_sanitize(form.achievement.data or ""),
            )
            db.session.add(p)
            db.session.commit()
            flash("Project added.", "success")
            return redirect(url_for("projects_index"))
        return render_template_string(PROJECT_FORM_TEMPLATE, form=form, project=None)

    @app.route("/projects/<int:project_id>/edit", methods=["GET", "POST"])
    @login_required
    def projects_edit(project_id):
        p = Project.query.get_or_404(project_id)
        if p.user_id != current_user.id:
            abort(403)
        form = ProjectForm(obj=p)
        if form.validate_on_submit():
            p.title = _sanitize(form.title.data)
            p.project_type = form.project_type.data
            p.start_date = _sanitize(form.start_date.data or "")
            p.end_date = _sanitize(form.end_date.data or "")
            p.github_url = _sanitize(form.github_url.data or "")
            p.live_url = _sanitize(form.live_url.data or "")
            p.description = _sanitize(form.description.data or "")
            p.technologies = _sanitize(form.technologies.data or "")
            p.achievement = _sanitize(form.achievement.data or "")
            db.session.commit()
            flash("Project updated.", "success")
            return redirect(url_for("projects_index"))
        return render_template_string(PROJECT_FORM_TEMPLATE, form=form, project=p)

    @app.route("/projects/<int:project_id>/delete", methods=["POST"])
    @login_required
    def projects_delete(project_id):
        p = Project.query.get_or_404(project_id)
        if p.user_id != current_user.id:
            abort(403)
        db.session.delete(p)
        db.session.commit()
        flash("Project deleted.", "info")
        return redirect(url_for("projects_index"))

    # ──────────────────────────────────────────────────────────────────────────
    # Error handlers
    # ──────────────────────────────────────────────────────────────────────────

    @app.errorhandler(403)
    def forbidden(e):
        return render_template_string(
            BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="text-center mt-5">
  <h1 class="display-1">403</h1>
  <p class="lead">You don't have permission to access this resource.</p>
  <a href="{{ url_for('dashboard') }}" class="btn btn-primary">Go Home</a>
</div>
{% endblock %}
"""),
            title="403 Forbidden",
        ), 403

    @app.errorhandler(404)
    def not_found(e):
        return render_template_string(
            BASE_TEMPLATE.replace("{% block content %}{% endblock %}", """
{% block content %}
<div class="text-center mt-5">
  <h1 class="display-1">404</h1>
  <p class="lead">Page not found.</p>
  <a href="{{ url_for('dashboard') }}" class="btn btn-primary">Go Home</a>
</div>
{% endblock %}
"""),
            title="404 Not Found",
        ), 404

    # V2+ routes registered from v2_features module (see bottom of file)


# =============================================================================
# Application entry point
# =============================================================================

app = create_app()

# ── Register V2+ features after templates are defined ─────────────────────────
_v2_ctx["BASE_TEMPLATE"] = BASE_TEMPLATE
_v2_ctx["csrf_hidden"] = _csrf_hidden
from v2_features import register_v2
register_v2(app, _v2_ctx)

if _v2_ctx.get("jobs") and not scheduler.get_job("weekly_reviews"):
    scheduler.add_job(func=_v2_ctx["jobs"]["weekly_reviews"], trigger="cron", day_of_week="sun", hour=23, minute=55, id="weekly_reviews")
    scheduler.add_job(func=_v2_ctx["jobs"]["mood_insights"], trigger="cron", day_of_week="mon", hour=6, minute=0, id="mood_insights")
    scheduler.add_job(func=_v2_ctx["jobs"]["email_digest"], trigger="cron", day_of_week="sun", hour=8, minute=0, id="email_digest")
    scheduler.add_job(func=_v2_ctx["jobs"]["recurring_tasks"], trigger="cron", hour=0, minute=5, id="recurring_tasks")

if __name__ == "__main__":
    app.run(debug=True)